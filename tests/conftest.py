"""
Konfiguracja pytest dla testów E2E Scenarzysta.
"""
import json
import os
import signal
import subprocess
import sys
import time
from pathlib import Path

import pytest
from playwright.sync_api import sync_playwright
import requests


# Bazowy URL aplikacji
BASE_URL = "http://127.0.0.1:5000"

# Ścieżka do plików testowych
TEST_FILES_DIR = Path(__file__).parent / "test_files"


@pytest.fixture(scope="session")
def test_files_dir():
    """Zwraca ścieżkę do katalogu z plikami testowymi."""
    TEST_FILES_DIR.mkdir(exist_ok=True)
    return TEST_FILES_DIR


@pytest.fixture(scope="session")
def sample_docx_path(test_files_dir):
    """Tworzy mały plik DOCX do testów: 3 zdania + 1 grafika.

    Dzięki temu przetwarzanie w testach E2E jest szybkie, ale nadal sprawdza
    ścieżkę z analizą obrazów.
    """
    from docx import Document

    doc_path = test_files_dir / "test_dokumentacja_mala.docx"
    image_path = test_files_dir / "test_obrazek.png"

    # Pobierz prostą grafikę z sieci (tylko przy pierwszym uruchomieniu)
    if not image_path.exists():
        try:
            resp = requests.get(
                "https://via.placeholder.com/300x200.png?text=Scenarzysta+Test",
                timeout=10,
            )
            resp.raise_for_status()
            image_path.write_bytes(resp.content)
        except Exception as e:
            print(f"[TEST] Nie udało się pobrać obrazka testowego: {e}")

    if not doc_path.exists():
        doc = Document()
        doc.add_heading("Mini dokumentacja testowa", 0)
        doc.add_paragraph(
            "System testowy 'Scenarzysta' służy do generowania scenariuszy testowych na podstawie dokumentacji."
        )
        doc.add_paragraph(
            "Użytkownik wgrywa dokument, wybiera opcje przetwarzania i otrzymuje ścieżki oraz scenariusze."
        )
        doc.add_paragraph(
            "Na końcu system może wygenerować także testy automatyczne w wybranej technologii."
        )
        if image_path.exists():
            # Dodaj jedną prostą grafikę na końcu dokumentu
            doc.add_picture(str(image_path))
        doc.save(doc_path)

    return doc_path


@pytest.fixture(scope="session")
def sample_excel_scenarios_path(test_files_dir):
    """Tworzy przykładowy plik Excel ze scenariuszami manualnymi."""
    from openpyxl import Workbook
    
    excel_path = test_files_dir / "test_scenariusze_manualne.xlsx"
    
    if not excel_path.exists():
        wb = Workbook()
        ws = wb.active
        ws.title = "Scenariusze"
        
        # Nagłówki
        headers = ['ID', 'Typ', 'Ścieżka', 'Tytuł', 'Nr kroku', 'Akcja', 'Oczekiwany rezultat', 'Warunki wstępne', 'Sekcja']
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Scenariusz 1 - Logowanie
        test_data = [
            ['TC001', 'Pozytywny', 'Logowanie', 'Poprawne logowanie użytkownika', 1, 'Otwórz stronę logowania', 'Strona logowania się wyświetla', 'Brak', 'Logowanie'],
            ['TC001', 'Pozytywny', 'Logowanie', 'Poprawne logowanie użytkownika', 2, 'Wpisz poprawny login', 'Login jest widoczny w polu', 'Brak', 'Logowanie'],
            ['TC001', 'Pozytywny', 'Logowanie', 'Poprawne logowanie użytkownika', 3, 'Wpisz poprawne hasło', 'Hasło jest zamaskowane', 'Brak', 'Logowanie'],
            ['TC001', 'Pozytywny', 'Logowanie', 'Poprawne logowanie użytkownika', 4, 'Kliknij przycisk Zaloguj', 'Użytkownik jest zalogowany', 'Brak', 'Logowanie'],
            ['TC002', 'Negatywny', 'Logowanie', 'Logowanie z błędnym hasłem', 1, 'Otwórz stronę logowania', 'Strona logowania się wyświetla', 'Brak', 'Logowanie'],
            ['TC002', 'Negatywny', 'Logowanie', 'Logowanie z błędnym hasłem', 2, 'Wpisz poprawny login', 'Login jest widoczny', 'Brak', 'Logowanie'],
            ['TC002', 'Negatywny', 'Logowanie', 'Logowanie z błędnym hasłem', 3, 'Wpisz błędne hasło', 'Hasło jest zamaskowane', 'Brak', 'Logowanie'],
            ['TC002', 'Negatywny', 'Logowanie', 'Logowanie z błędnym hasłem', 4, 'Kliknij przycisk Zaloguj', 'Wyświetla się komunikat o błędzie', 'Brak', 'Logowanie'],
        ]
        
        for row_idx, row_data in enumerate(test_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        wb.save(excel_path)
    
    return excel_path


@pytest.fixture(scope="session")
def sample_example_template_path(test_files_dir):
    """Tworzy przykładowy plik JSON z dokumentacją i scenariuszami.

    Plik jest w formacie zgodnym z example_template.json, ale zawiera
    bardzo krótki, testowy przykład.
    """
    template_path = test_files_dir / "test_example_template.json"

    if not template_path.exists():
        data = {
            "_instrukcja": "Przykładowy szablon testowy dla Scenarzysty v0.6.",
            "example_documentation": (
                "System TEST v0.6 obsługuje logowanie użytkownika, reset hasła "
                "oraz podgląd profilu. Dokumentacja ma tylko kilka zdań, aby "
                "przetwarzanie w testach było szybkie."
            ),
            "example_scenarios": [
                {
                    "scenario_title": "Logowanie z poprawnymi danymi (TEST)",
                    "path_type": "positive",
                    "test_path": "Logowanie TEST v0.6",
                    "prerequisites": [
                        "użytkownik ma aktywne konto",
                        "użytkownik zna login i hasło",
                    ],
                    "steps": [
                        {
                            "step_number": 1,
                            "action": "Otwórz stronę logowania TEST.",
                            "expected_result": "Formularz logowania jest widoczny.",
                        },
                        {
                            "step_number": 2,
                            "action": "Wpisz poprawny login i hasło.",
                            "expected_result": "Dane są widoczne w polach formularza.",
                        },
                        {
                            "step_number": 3,
                            "action": "Kliknij przycisk 'Zaloguj'.",
                            "expected_result": "System przekierowuje na stronę główną TEST.",
                        },
                    ],
                }
            ],
        }
        template_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    return template_path


@pytest.fixture(scope="session")
def app_server():
    """
    Uruchamia serwer Flask przed testami i zamyka po zakończeniu.
    Jeśli serwer już działa, używa istniejącej instancji.
    """
    # Sprawdź czy serwer już działa
    try:
        response = requests.get(f"{BASE_URL}/api/health", timeout=2)
        if response.status_code == 200:
            print("Serwer już działa - używam istniejącej instancji")
            yield BASE_URL
            return
    except:
        pass
    
    # Uruchom serwer
    project_dir = Path(__file__).parent.parent
    env = os.environ.copy()
    env['FLASK_ENV'] = 'testing'
    # Domyślny model Ollama na potrzeby testów E2E (jeśli nie ustawiono inaczej)
    env.setdefault('OLLAMA_MODEL', 'gemma3:12B')
    
    process = subprocess.Popen(
        [sys.executable, 'app.py'],
        cwd=project_dir,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        preexec_fn=os.setsid
    )
    
    # Poczekaj na uruchomienie serwera
    max_wait = 30
    for _ in range(max_wait):
        try:
            response = requests.get(f"{BASE_URL}/api/health", timeout=1)
            if response.status_code == 200:
                break
        except Exception:
            time.sleep(1)
    else:
        process.terminate()
        raise RuntimeError("Nie udało się uruchomić serwera")
    
    yield BASE_URL
    
    # Zamknij serwer
    os.killpg(os.getpgid(process.pid), signal.SIGTERM)
    process.wait()


@pytest.fixture(scope="function")
def page(app_server):
    """Tworzy nową stronę Playwright dla każdego testu."""
    with sync_playwright() as p:
        headless_env = os.environ.get("HEADLESS", "true").lower()
        headless = headless_env not in ("false", "0", "no")
        slow_mo_str = os.environ.get("PLAYWRIGHT_SLOWMO", "0") or "0"
        try:
            slow_mo = int(slow_mo_str)
        except ValueError:
            slow_mo = 0

        browser = p.chromium.launch(headless=headless, slow_mo=slow_mo)
        context = browser.new_context()
        page = context.new_page()
        page.goto(app_server)
        yield page
        context.close()
        browser.close()
