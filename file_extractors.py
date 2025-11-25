"""
Moduł ekstrakcji tekstu z różnych formatów plików.
Obsługuje: DOCX, PDF, XLSX/XLS, TXT
"""
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass


@dataclass
class ExtractedContent:
    """Wynik ekstrakcji z pliku."""
    text: str
    images: List[Dict]  # Lista obrazów z metadanymi
    tables: List[Dict]  # Lista tabel
    metadata: Dict  # Metadane dokumentu
    page_count: int
    source_type: str  # docx, pdf, xlsx, txt


class FileExtractor:
    """Ekstraktor tekstu z różnych formatów plików."""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.xlsx', '.xls', '.txt'}
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Sprawdza czy format pliku jest obsługiwany."""
        ext = Path(filename).suffix.lower()
        return ext in FileExtractor.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Zwraca typ pliku na podstawie rozszerzenia."""
        ext = Path(filename).suffix.lower()
        return ext.lstrip('.')
    
    def extract(self, file_path: str, output_dir: str = None) -> ExtractedContent:
        """
        Ekstrahuje tekst z pliku.
        
        Args:
            file_path: Ścieżka do pliku
            output_dir: Katalog na obrazy (dla docx/pdf)
            
        Returns:
            ExtractedContent z tekstem i metadanymi
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.docx':
            return self._extract_docx(file_path, output_dir)
        elif ext == '.pdf':
            return self._extract_pdf(file_path, output_dir)
        elif ext in ['.xlsx', '.xls']:
            return self._extract_excel(file_path)
        elif ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Nieobsługiwany format pliku: {ext}")
    
    def _extract_docx(self, file_path: str, output_dir: str = None) -> ExtractedContent:
        """Ekstrahuje tekst z pliku DOCX."""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ImportError("Zainstaluj python-docx: pip install python-docx")
        
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Nie można otworzyć pliku DOCX: {file_path}")
        
        # Ekstrahuj tekst
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Ekstrahuj tabele
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'index': table_idx,
                'data': table_data,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0
            })
        
        # Ekstrahuj obrazy (jeśli podano output_dir)
        images = []
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        image_ext = rel.target_part.content_type.split('/')[-1]
                        image_name = f"image_{len(images)}.{image_ext}"
                        image_path = output_path / image_name
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_data)
                        
                        images.append({
                            'path': str(image_path),
                            'name': image_name,
                            'type': image_ext
                        })
                    except Exception as e:
                        print(f"  Błąd ekstrakcji obrazu: {e}")
        
        full_text = "\n\n".join(paragraphs)
        
        # Dodaj tabele do tekstu
        for table in tables:
            table_text = f"\n[TABELA {table['index'] + 1}]\n"
            for row in table['data']:
                table_text += " | ".join(row) + "\n"
            full_text += table_text
        
        return ExtractedContent(
            text=full_text,
            images=images,
            tables=tables,
            metadata={'filename': Path(file_path).name},
            page_count=len(doc.sections),
            source_type='docx'
        )
    
    def _extract_pdf(self, file_path: str, output_dir: str = None) -> ExtractedContent:
        """Ekstrahuje tekst z pliku PDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            # Fallback do pdfplumber
            try:
                import pdfplumber
                return self._extract_pdf_pdfplumber(file_path, output_dir)
            except ImportError:
                raise ImportError("Zainstaluj PyMuPDF lub pdfplumber: pip install pymupdf pdfplumber")
        
        doc = fitz.open(file_path)
        
        all_text = []
        images = []
        tables = []
        
        for page_num, page in enumerate(doc):
            # Tekst
            text = page.get_text()
            if text.strip():
                all_text.append(f"[Strona {page_num + 1}]\n{text}")
            
            # Obrazy (jeśli podano output_dir)
            if output_dir:
                output_path = Path(output_dir)
                output_path.mkdir(parents=True, exist_ok=True)
                
                image_list = page.get_images()
                for img_idx, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        image_name = f"page{page_num + 1}_img{img_idx}.{image_ext}"
                        image_path = output_path / image_name
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_bytes)
                        
                        images.append({
                            'path': str(image_path),
                            'name': image_name,
                            'type': image_ext,
                            'page': page_num + 1
                        })
                    except Exception as e:
                        print(f"  Błąd ekstrakcji obrazu ze strony {page_num + 1}: {e}")
        
        doc.close()
        
        return ExtractedContent(
            text="\n\n".join(all_text),
            images=images,
            tables=tables,
            metadata={'filename': Path(file_path).name, 'pages': len(all_text)},
            page_count=len(all_text),
            source_type='pdf'
        )
    
    def _extract_pdf_pdfplumber(self, file_path: str, output_dir: str = None) -> ExtractedContent:
        """Fallback ekstrakcji PDF przez pdfplumber."""
        import pdfplumber
        
        all_text = []
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Tekst
                text = page.extract_text()
                if text:
                    all_text.append(f"[Strona {page_num + 1}]\n{text}")
                
                # Tabele
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    tables.append({
                        'index': len(tables),
                        'page': page_num + 1,
                        'data': table,
                        'rows': len(table) if table else 0
                    })
        
        # Dodaj tabele do tekstu
        full_text = "\n\n".join(all_text)
        for table in tables:
            table_text = f"\n[TABELA ze strony {table['page']}]\n"
            for row in table['data']:
                if row:
                    table_text += " | ".join(str(cell) if cell else '' for cell in row) + "\n"
            full_text += table_text
        
        return ExtractedContent(
            text=full_text,
            images=[],  # pdfplumber nie ekstrahuje obrazów łatwo
            tables=tables,
            metadata={'filename': Path(file_path).name},
            page_count=len(all_text),
            source_type='pdf'
        )
    
    def _extract_excel(self, file_path: str) -> ExtractedContent:
        """Ekstrahuje dane z pliku Excel."""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("Zainstaluj openpyxl: pip install openpyxl")
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        all_text = []
        tables = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            # Nagłówek arkusza
            all_text.append(f"\n[ARKUSZ: {sheet_name}]")
            
            # Pobierz dane jako tabelę
            table_data = []
            for row in sheet.iter_rows():
                row_data = []
                for cell in row:
                    value = cell.value
                    if value is not None:
                        row_data.append(str(value))
                    else:
                        row_data.append('')
                if any(row_data):  # Pomijaj puste wiersze
                    table_data.append(row_data)
            
            if table_data:
                tables.append({
                    'index': len(tables),
                    'sheet': sheet_name,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': max(len(row) for row in table_data) if table_data else 0
                })
                
                # Dodaj do tekstu
                for row in table_data:
                    all_text.append(" | ".join(row))
        
        wb.close()
        
        return ExtractedContent(
            text="\n".join(all_text),
            images=[],
            tables=tables,
            metadata={
                'filename': Path(file_path).name,
                'sheets': wb.sheetnames
            },
            page_count=len(wb.sheetnames),
            source_type='xlsx'
        )
    
    def _extract_txt(self, file_path: str) -> ExtractedContent:
        """Ekstrahuje tekst z pliku TXT."""
        encodings = ['utf-8', 'cp1250', 'iso-8859-2', 'latin-1']
        
        text = None
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError(f"Nie można odczytać pliku TXT: {file_path}")
        
        # Policz "strony" (każde 3000 znaków = 1 strona)
        page_count = max(1, len(text) // 3000)
        
        return ExtractedContent(
            text=text,
            images=[],
            tables=[],
            metadata={'filename': Path(file_path).name, 'encoding': encoding},
            page_count=page_count,
            source_type='txt'
        )


def extract_file(file_path: str, output_dir: str = None) -> ExtractedContent:
    """
    Funkcja pomocnicza do ekstrakcji pliku.
    
    Args:
        file_path: Ścieżka do pliku
        output_dir: Katalog na obrazy
        
    Returns:
        ExtractedContent
    """
    extractor = FileExtractor()
    return extractor.extract(file_path, output_dir)
