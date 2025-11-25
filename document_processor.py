"""
Procesor dokumentów - ekstrakcja i przetwarzanie dokumentów.
Obsługuje formaty: DOCX, PDF, XLSX, TXT
Dane są przetwarzane tylko dla konkretnego przypadku, bez trwałego przechowywania w RAG.
"""
import zipfile
import os
import re
import base64
import requests
from pathlib import Path
from typing import Dict, List, Optional, Any
import json
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

# Import ekstraktora dla innych formatów
try:
    from file_extractors import FileExtractor, extract_file
    FILE_EXTRACTORS_AVAILABLE = True
except ImportError:
    FILE_EXTRACTORS_AVAILABLE = False


class ContextLengthError(Exception):
    """Błąd przekroczenia limitu kontekstu/tokenów modelu."""
    pass


class DocumentProcessor:
    """Procesor dokumentów z ekstrakcją multimodalną."""
    
    def __init__(self, ollama_url: str = "http://localhost:11434", ollama_model: str = "gemma2:2b"):
        """
        Inicjalizuje procesor dokumentów.
        
        Args:
            ollama_url: URL serwera Ollama (domyślnie localhost:11434)
            ollama_model: Nazwa modelu wizyjnego Ollama (domyślnie gemma2:2b, można zmienić na gemma3)
        """
        self.ollama_url = ollama_url
        self.ollama_model = ollama_model
        
        # Śledzenie postępu dla dynamicznej estymacji czasu
        self.processing_stats = {
            'total_chunks': 0,
            'processed_chunks': 0,
            'chunk_times': [],  # Lista czasów przetwarzania chunków
            'start_time': None,
            'current_stage': 0,
            'total_stages': 3
        }
    
    def reset_processing_stats(self):
        """Resetuje statystyki przetwarzania."""
        self.processing_stats = {
            'total_chunks': 0,
            'processed_chunks': 0,
            'chunk_times': [],
            'start_time': None,
            'current_stage': 0,
            'total_stages': 3
        }
    
    def get_dynamic_eta(self) -> Optional[float]:
        """
        Oblicza dynamiczny ETA na podstawie rzeczywistego postępu.
        
        Returns:
            Szacowany pozostały czas w sekundach lub None
        """
        import time
        stats = self.processing_stats
        
        if not stats['chunk_times'] or stats['total_chunks'] == 0:
            return None
        
        # Średni czas na chunk
        avg_chunk_time = sum(stats['chunk_times']) / len(stats['chunk_times'])
        
        # Pozostałe chunki w bieżącym etapie
        remaining_chunks = stats['total_chunks'] - stats['processed_chunks']
        
        # Pozostałe etapy (każdy etap ma podobną liczbę chunków)
        remaining_stages = stats['total_stages'] - stats['current_stage']
        
        # Estymacja: pozostałe chunki * średni czas + pozostałe etapy * (średni czas * średnia liczba chunków)
        eta = remaining_chunks * avg_chunk_time
        if remaining_stages > 0 and stats['total_chunks'] > 0:
            eta += remaining_stages * stats['total_chunks'] * avg_chunk_time
        
        return eta
    
    def _filter_header_footer_images(self, images: List[Dict]) -> List[Dict]:
        """
        Filtruje obrazy z nagłówka i stopki (zazwyczaj loga).
        
        Heurystyka:
        - Obrazy o nazwie zawierającej 'header', 'footer', 'logo'
        - Bardzo małe obrazy (< 50x50 px) - często ikony
        - Obrazy powtarzające się (ten sam hash) - loga na każdej stronie
        
        Args:
            images: Lista słowników z informacjami o obrazach
            
        Returns:
            Przefiltrowana lista obrazów
        """
        filtered = []
        seen_sizes = {}  # Rozmiar -> liczba wystąpień
        
        for img in images:
            filename = img.get('filename', '').lower()
            path = img.get('path', '')
            
            # Pomiń obrazy z nazwą sugerującą logo/nagłówek/stopkę
            skip_keywords = ['header', 'footer', 'logo', 'banner', 'watermark']
            if any(kw in filename for kw in skip_keywords):
                print(f"  Pomijam obraz (nazwa): {filename}")
                continue
            
            # Sprawdź rozmiar obrazu
            try:
                with Image.open(path) as pil_img:
                    width, height = pil_img.size
                    
                    # Pomiń bardzo małe obrazy (ikony, bullet points)
                    if width < 50 or height < 50:
                        print(f"  Pomijam obraz (za mały {width}x{height}): {filename}")
                        continue
                    
                    # Śledź rozmiary do wykrycia powtarzających się logo
                    size_key = f"{width}x{height}"
                    seen_sizes[size_key] = seen_sizes.get(size_key, 0) + 1
                    
            except Exception:
                # Jeśli nie można otworzyć obrazu, zachowaj go
                pass
            
            filtered.append(img)
        
        # Usuń obrazy o rozmiarze powtarzającym się >3 razy (prawdopodobnie loga)
        final_filtered = []
        for img in filtered:
            try:
                with Image.open(img.get('path', '')) as pil_img:
                    size_key = f"{pil_img.size[0]}x{pil_img.size[1]}"
                    if seen_sizes.get(size_key, 0) > 3:
                        print(f"  Pomijam obraz (powtarzający się): {img.get('filename')}")
                        continue
            except Exception:
                pass
            final_filtered.append(img)
        
        print(f"  Obrazy po filtracji: {len(final_filtered)}/{len(images)}")
        return final_filtered
    
    def extract_from_file(self, file_path: str, output_dir: str) -> Dict:
        """
        Ekstrahuje tekst z pliku (obsługuje różne formaty).
        
        Args:
            file_path: Ścieżka do pliku
            output_dir: Katalog wyjściowy dla ekstrahowanych danych
            
        Returns:
            Słownik z ekstrahowanymi danymi
        """
        ext = Path(file_path).suffix.lower()
        
        # Dla DOCX używaj natywnej metody (lepsza obsługa obrazów)
        if ext == '.docx':
            return self.extract_from_docx(file_path, output_dir)
        
        # Dla innych formatów użyj file_extractors
        if not FILE_EXTRACTORS_AVAILABLE:
            raise ImportError(f"Brak modułu file_extractors do obsługi plików {ext}")
        
        print(f"  Ekstrakcja z pliku {ext}: {Path(file_path).name}")
        
        extractor = FileExtractor()
        extracted = extractor.extract(file_path, output_dir)
        
        # Konwertuj na format zgodny z resztą systemu
        extracted_data = {
            'text': [{
                'section': 'Cały dokument',
                'content': extracted.text,
                'paragraph_count': extracted.text.count('\n') + 1,
                'image_placeholders': []
            }],
            'images': [
                {
                    'filename': img['name'],
                    'path': img['path'],
                    'original_path': img.get('path', '')
                }
                for img in extracted.images
            ],
            'metadata': {
                'filename': Path(file_path).name,
                'total_images': len(extracted.images),
                'total_sections': 1,
                'total_paragraphs': extracted.text.count('\n') + 1,
                'source_type': extracted.source_type,
                'page_count': extracted.page_count
            }
        }
        
        # Dla Excel dodaj informacje o tabelach
        if extracted.tables:
            extracted_data['metadata']['tables'] = len(extracted.tables)
        
        print(f"  Wyekstrahowano {len(extracted_data['text'][0]['content'])} znaków tekstu")
        
        return extracted_data
    
    def extract_from_docx(self, docx_path: str, output_dir: str) -> Dict:
        """
        Ekstrahuje tekst i obrazy z pliku .docx.
        
        Args:
            docx_path: Ścieżka do pliku .docx
            output_dir: Katalog wyjściowy dla ekstrahowanych danych
            
        Returns:
            Słownik z ekstrahowanymi danymi
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        images_dir = output_path / "images"
        images_dir.mkdir(exist_ok=True)
        
        extracted_data = {
            'text': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            # Otwórz dokument używając python-docx
            doc = Document(docx_path)
            
            # Ekstrahuj obrazy z archiwum ZIP
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                for file_info in zip_ref.namelist():
                    if file_info.startswith('word/media/'):
                        image_data = zip_ref.read(file_info)
                        image_filename = os.path.basename(file_info)
                        image_path = images_dir / image_filename
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        extracted_data['images'].append({
                            'filename': image_filename,
                            'path': str(image_path),
                            'original_path': file_info
                        })
                
            # Mapowanie obrazów do ich nazw plików
            image_filename_map = {img['filename']: img for img in extracted_data['images']}
            
            # Ekstrahuj tekst z dokumentu i śledź pozycje obrazów
            current_section = None
            section_content = []
            paragraph_count = 0
            image_placeholders = {}  # Mapa: pozycja w tekście -> informacja o obrazie
            
            def iter_block_items(parent):
                """Iteruje przez wszystkie bloki (paragrafy i tabele) w dokumencie."""
                if isinstance(parent, DocumentType):
                    parent_elm = parent.element.body
                elif isinstance(parent, Table):
                    parent_elm = parent._tbl
                else:
                    return
                
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)
            
            # Przetwarzaj wszystkie paragrafy i tabele
            block_index = 0
            for block in iter_block_items(doc):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    
                    # Sprawdź czy paragraf zawiera obrazy
                    paragraph_images = []
                    try:
                        for run in block.runs:
                            # Sprawdź czy run zawiera obraz
                            if run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                                # Znajdź wszystkie relacje obrazów w tym paragrafie
                                for drawing in run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                                    rId = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if rId:
                                        # Znajdź obraz w relacjach dokumentu
                                        try:
                                            rel = doc.part.rels[rId]
                                            if rel and 'image' in rel.target_ref:
                                                # Znajdź nazwę pliku obrazu
                                                image_path = rel.target_ref
                                                image_filename = os.path.basename(image_path)
                                                if image_filename in image_filename_map:
                                                    paragraph_images.append(image_filename)
                                        except KeyError:
                                            # Relacja nie istnieje - pomiń
                                            pass
                    except Exception as e:
                        # Jeśli wystąpi błąd podczas wykrywania obrazów, kontynuuj
                        pass
                    
                    # Jeśli paragraf zawiera obrazy, dodaj placeholder
                    if paragraph_images:
                        for img_filename in paragraph_images:
                            placeholder = f"__IMAGE_PLACEHOLDER_{img_filename}__"
                            if text:
                                text = f"{text}\n{placeholder}"
                            else:
                                text = placeholder
                            
                            # Zapisz informację o obrazie i jego pozycji
                            image_placeholders[placeholder] = {
                                'filename': img_filename,
                                'section': current_section or f'Sekcja {len(extracted_data["text"]) + 1}',
                                'position': len(section_content),
                                'image_info': image_filename_map[img_filename]
                            }
                    
                    if not text:
                        continue
                    
                    # Sprawdź czy to nagłówek (styl zaczynający się od "Heading")
                    is_heading = False
                    heading_level = 0
                    if block.style and block.style.name.startswith('Heading'):
                        is_heading = True
                        # Wyciągnij poziom nagłówka (np. "Heading 1" -> 1)
                        match = re.search(r'Heading\s+(\d+)', block.style.name)
                        if match:
                            heading_level = int(match.group(1))
                    
                    # Jeśli to nagłówek, zapisz poprzednią sekcję i rozpocznij nową
                    if is_heading:
                        if current_section and section_content:
                            extracted_data['text'].append({
                                'section': current_section,
                                'content': '\n'.join(section_content),
                                'paragraph_count': len(section_content),
                                'image_placeholders': [ph for ph in image_placeholders.values() 
                                                       if ph['section'] == current_section]
                            })
                        current_section = text
                        section_content = []
                    else:
                        # Dodaj tekst do bieżącej sekcji
                        if not current_section:
                            current_section = f'Sekcja {len(extracted_data["text"]) + 1}'
                        section_content.append(text)
                        paragraph_count += 1
                
                elif isinstance(block, Table):
                    # Ekstrahuj tekst z tabel
                    table_text = []
                    for row in block.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    
                    if table_text:
                        if not current_section:
                            current_section = f'Tabela {len(extracted_data["text"]) + 1}'
                        section_content.append('\n'.join(table_text))
                
                block_index += 1
            
            # Zapisz ostatnią sekcję
            if current_section and section_content:
                extracted_data['text'].append({
                    'section': current_section,
                    'content': '\n'.join(section_content),
                    'paragraph_count': len(section_content),
                    'image_placeholders': [ph for ph in image_placeholders.values() 
                                           if ph['section'] == current_section]
                })
            
            # Jeśli nie znaleziono żadnych sekcji, użyj całego tekstu
            if not extracted_data['text']:
                full_text = []
                full_text_placeholders = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    
                    # Sprawdź czy paragraf zawiera obrazy (fallback)
                    try:
                        for run in para.runs:
                            if run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                                for drawing in run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                                    rId = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if rId:
                                        try:
                                            rel = doc.part.rels[rId]
                                            if rel and 'image' in rel.target_ref:
                                                image_filename = os.path.basename(rel.target_ref)
                                                if image_filename in image_filename_map:
                                                    placeholder = f"__IMAGE_PLACEHOLDER_{image_filename}__"
                                                    if text:
                                                        text = f"{text}\n{placeholder}"
                                                    else:
                                                        text = placeholder
                                                    full_text_placeholders.append({
                                                        'filename': image_filename,
                                                        'section': 'Cały dokument',
                                                        'position': len(full_text),
                                                        'image_info': image_filename_map[image_filename]
                                                    })
                                        except KeyError:
                                            pass
                    except Exception:
                        pass
                    
                    if text:
                        full_text.append(text)
                
                if full_text:
                    extracted_data['text'].append({
                        'section': 'Cały dokument',
                        'content': '\n'.join(full_text),
                        'paragraph_count': len(full_text),
                        'image_placeholders': full_text_placeholders
                    })
                    
                    if full_text_placeholders:
                        print(f"  Wstawiono {len(full_text_placeholders)} placeholderów obrazów do tekstu")
            
            # Policz wszystkie placeholdery obrazów
            total_placeholders = sum(
                len(section.get('image_placeholders', [])) 
                for section in extracted_data['text']
            )
            
            extracted_data['metadata'] = {
                'filename': os.path.basename(docx_path),
                'total_images': len(extracted_data['images']),
                'total_sections': len(extracted_data['text']),
                'total_paragraphs': paragraph_count,
                'total_image_placeholders': total_placeholders,
                'extraction_time': str(Path(docx_path).stat().st_mtime)
            }
            
            if total_placeholders > 0:
                print(f"  Wykryto {total_placeholders} obrazów powiązanych z tekstem")
        
        except Exception as e:
            raise Exception(f"Błąd podczas ekstrakcji z .docx: {str(e)}")
        
        return extracted_data
    
    def analyze_image_with_ollama(self, image_path: str) -> Optional[str]:
        """
        Analizuje obraz za pomocą Ollama z modelem wizyjnym.
        
        Args:
            image_path: Ścieżka do obrazu
            
        Returns:
            Opis obrazu lub None w przypadku błędu
        """
        try:
            # Sprawdź czy obraz istnieje
            if not os.path.exists(image_path):
                return None
            
            # Wczytaj i zakoduj obraz w base64
            with open(image_path, 'rb') as img_file:
                image_data = base64.b64encode(img_file.read()).decode('utf-8')
            
            # Określ typ MIME obrazu
            img_ext = Path(image_path).suffix.lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.bmp': 'image/bmp',
                '.webp': 'image/webp'
            }
            mime_type = mime_types.get(img_ext, 'image/jpeg')
            
            # Przygotuj prompt dla modelu wizyjnego
            prompt = """Przeanalizuj ten obraz z dokumentacji technicznej. 
            Opisz co widzisz na obrazie, zwracając szczególną uwagę na:
            - Elementy interfejsu użytkownika (przyciski, pola, menu)
            - Diagramy, schematy, wykresy
            - Tekst widoczny na obrazie
            - Funkcjonalności przedstawione na obrazie
            - Wszelkie informacje istotne dla testowania
            
            Odpowiedz po polsku, zwięźle i konkretnie."""
            
            # Wywołaj Ollama API dla modeli wizyjnych
            # Użyj /api/chat dla modeli wizyjnych (gemma2, gemma3, llava, itp.)
            api_url = f"{self.ollama_url}/api/chat"
            
            payload = {
                "model": self.ollama_model,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt,
                        "images": [image_data]
                    }
                ],
                "stream": False
            }
            
            response = requests.post(api_url, json=payload, timeout=120)  # Dłuższy timeout dla analizy obrazów
            
            if response.status_code == 200:
                result = response.json()
                # Dla /api/chat odpowiedź jest w message.content
                message = result.get('message', {})
                description = message.get('content', '').strip()
                # Fallback na starą strukturę
                if not description:
                    description = result.get('response', '').strip()
                return description if description else None
            elif response.status_code == 500:
                error_text = response.text.lower()
                # Wykryj błędy pamięci GPU
                if 'resource' in error_text or 'memory' in error_text or 'stopped' in error_text:
                    print(f"  ⚠️ Błąd pamięci GPU podczas analizy obrazu - pomijam obraz: {os.path.basename(image_path)}")
                    print(f"  💡 Wskazówka: Spróbuj zmniejszyć rozmiar obrazów lub użyć mniejszego modelu")
                    return None
                print(f"Błąd serwera Ollama (500): {response.text[:200]}")
                return None
            else:
                print(f"Błąd podczas analizy obrazu przez Ollama: {response.status_code} - {response.text}")
                return None
                
        except requests.exceptions.Timeout:
            print(f"  ⚠️ Timeout podczas analizy obrazu: {os.path.basename(image_path)} - pomijam")
            return None
        except requests.exceptions.RequestException as e:
            print(f"Błąd połączenia z Ollama: {e}")
            return None
        except Exception as e:
            print(f"Błąd podczas analizy obrazu {image_path}: {e}")
            return None
    
    def analyze_multimodal(self, extracted_data: Dict, processing_dir: Path, analyze_images: bool = False) -> Dict:
        """
        Analizuje ekstrahowane dane multimodalne.
        
        Args:
            extracted_data: Ekstrahowane dane
            processing_dir: Katalog przetwarzania
            analyze_images: Czy analizować obrazy przez LLM (domyślnie False - szybsze przetwarzanie)
            
        Returns:
            Przeanalizowane dane
        """
        analyzed_data = {
            'text_analysis': [],
            'image_analysis': [],
            'combined_insights': []
        }
        
        # Analiza obrazów przez Ollama (opcjonalna)
        image_descriptions = {}  # Mapa: filename -> description
        
        if analyze_images:
            # Filtruj obrazy z nagłówka/stopki
            images_to_analyze = self._filter_header_footer_images(extracted_data.get('images', []))
            
            print(f"  Analiza {len(images_to_analyze)} obrazów przez LLM...")
            for idx, image_item in enumerate(images_to_analyze, 1):
                image_path = image_item.get('path')
                if image_path and os.path.exists(image_path):
                    print(f"    Analizuję obraz {idx}/{len(images_to_analyze)}: {image_item['filename']}")
                    # Analizuj obraz przez Ollama
                    description = self.analyze_image_with_ollama(image_path)
                    
                    if description:
                        image_descriptions[image_item['filename']] = description
                        analyzed_data['image_analysis'].append({
                            'filename': image_item['filename'],
                            'description': description,
                            'ui_elements': [],
                            'text_from_image': description
                        })
                    else:
                        analyzed_data['image_analysis'].append({
                            'filename': image_item['filename'],
                            'description': f"Obraz {image_item['filename']} (analiza nie powiodła się)",
                            'ui_elements': [],
                            'text_from_image': ''
                        })
        else:
            print("  Pominięto analizę obrazów (opcja wyłączona - szybsze przetwarzanie)")
        
        # TERAZ: Wstaw opisy obrazów w odpowiednie miejsca w tekście PRZED analizą
        text_items_with_images = []
        for text_item in extracted_data.get('text', []):
            content = text_item.get('content', '')
            section = text_item.get('section', '')
            
            # Znajdź placeholdery obrazów w tekście i zastąp je opisami
            image_placeholders_list = text_item.get('image_placeholders', [])
            for image_placeholder_info in image_placeholders_list:
                filename = image_placeholder_info.get('filename', '')
                if filename in image_descriptions:
                    description = image_descriptions[filename]
                    # Zastąp placeholder opisem obrazu
                    placeholder_text = f"__IMAGE_PLACEHOLDER_{filename}__"
                    image_description_text = f"\n[OPIS OBRAZU: {description}]\n"
                    content = content.replace(placeholder_text, image_description_text)
            
            # Zaktualizuj zawartość z opisami obrazów
            text_item['content'] = content
            text_items_with_images.append(text_item)
        
        # Analiza tekstu (z opisami obrazów już wstawionymi) - identyfikacja wymagań, funkcjonalności, scenariuszy
        all_requirements = []
        all_functionalities = []
        all_test_scenarios = []
        
        for text_item in text_items_with_images:
            section = text_item.get('section', '')
            content = text_item.get('content', '')
            
            # Identyfikuj wymagania (szukaj wzorców jak "wymaganie", "REQ", "requirement", "musi", "powinien")
            requirements = self._extract_requirements(content, section)
            all_requirements.extend(requirements)
            
            # Identyfikuj funkcjonalności (szukaj wzorców jak "funkcja", "funkcjonalność", "feature", "akcja")
            functionalities = self._extract_functionalities(content, section)
            all_functionalities.extend(functionalities)
            
            # Identyfikuj potencjalne scenariusze testowe (szukaj wzorców jak "scenariusz", "przypadek", "test case")
            test_scenarios = self._extract_test_scenarios(content, section)
            all_test_scenarios.extend(test_scenarios)
            
            analyzed_data['text_analysis'].append({
                'section': section,
                'content': content,
                'requirements': requirements,
                'functionalities': functionalities,
                'test_scenarios': test_scenarios
            })
        
        # Połączone wnioski - wszystkie znalezione wymagania, funkcjonalności i scenariusze
        combined_insights = []
        
        # Dodaj wymagania
        for req in all_requirements:
            combined_insights.append({
                'type': 'requirement',
                'description': req,
                'source': 'text',
                'confidence': 0.8
            })
        
        # Dodaj funkcjonalności
        for func in all_functionalities:
            combined_insights.append({
                'type': 'functionality',
                'description': func,
                'source': 'text',
                'confidence': 0.7
            })
        
        # Dodaj scenariusze testowe
        for scenario in all_test_scenarios:
            combined_insights.append({
                'type': 'test_scenario',
                'description': scenario,
                'source': 'text',
                'confidence': 0.75
            })
        
        # Jeśli nie znaleziono żadnych wniosków, użyj sekcji jako wymagań
        if not combined_insights:
            for text_item in extracted_data.get('text', []):
                section = text_item.get('section', '')
                content = text_item.get('content', '')
                # Użyj pierwszych 200 znaków jako opis
                short_content = content[:200] + ('...' if len(content) > 200 else '')
                combined_insights.append({
                    'type': 'requirement',
                    'description': f"{section}: {short_content}",
                    'source': 'text',
                    'confidence': 0.6
                })
        
        analyzed_data['combined_insights'] = combined_insights
        
        # Zwróć extracted_data z dodanymi opisami obrazów (dla kompatybilności z nowymi metodami)
        # Dodaj tekst z opisami obrazów z powrotem do extracted_data
        extracted_data['text'] = text_items_with_images
        extracted_data['image_descriptions'] = image_descriptions
        
        return extracted_data
    
    def _extract_requirements(self, content: str, section: str) -> List[str]:
        """Ekstrahuje wymagania z tekstu."""
        requirements = []
        
        # Wzorce do identyfikacji wymagań
        patterns = [
            r'(?:wymaganie|requirement|REQ)[\s:]+([^\.\n]+)',
            r'(?:system|aplikacja|moduł)[\s]+(?:musi|powinien|powinna|powinno)[\s]+([^\.\n]+)',
            r'(?:funkcjonalność|funkcja)[\s:]+([^\.\n]+)',
            r'([A-Z][^\.\n]{20,200}(?:musi|powinien|powinna|powinno|wymaga|obsługuje)[^\.\n]{10,200})',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                req_text = match.group(1 if match.lastindex else 0).strip()
                if len(req_text) > 10:  # Minimum długość wymagania
                    requirements.append(req_text)
        
        # Jeśli nie znaleziono wzorców, użyj całych zdań zawierających kluczowe słowa
        if not requirements:
            sentences = re.split(r'[\.!?]\s+', content)
            keywords = ['wymaganie', 'requirement', 'musi', 'powinien', 'funkcja', 'funkcjonalność']
            for sentence in sentences:
                if any(keyword.lower() in sentence.lower() for keyword in keywords):
                    if len(sentence.strip()) > 15:
                        requirements.append(sentence.strip())
        
        return requirements[:20]  # Maksymalnie 20 wymagań na sekcję
    
    def _extract_functionalities(self, content: str, section: str) -> List[str]:
        """Ekstrahuje funkcjonalności z tekstu."""
        functionalities = []
        
        patterns = [
            r'(?:funkcja|funkcjonalność|feature)[\s:]+([^\.\n]+)',
            r'(?:umożliwia|obsługuje|realizuje)[\s]+([^\.\n]+)',
            r'(?:użytkownik|użytkownicy)[\s]+(?:może|mogą)[\s]+([^\.\n]+)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                func_text = match.group(1 if match.lastindex else 0).strip()
                if len(func_text) > 10:
                    functionalities.append(func_text)
        
        return functionalities[:20]
    
    def _extract_test_scenarios(self, content: str, section: str) -> List[str]:
        """Ekstrahuje scenariusze testowe z tekstu."""
        scenarios = []
        
        patterns = [
            r'(?:scenariusz|przypadek testowy|test case)[\s:]+([^\.\n]+)',
            r'(?:gdy|jeśli|kiedy)[\s]+([^\.\n]{20,200}(?:wtedy|następnie|powinien)[^\.\n]{10,200})',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                scenario_text = match.group(1 if match.lastindex else 0).strip()
                if len(scenario_text) > 15:
                    scenarios.append(scenario_text)
        
        return scenarios[:20]
    
    def generate_test_scenarios(self, analyzed_data: Dict) -> List[Dict]:
        """
        Generuje scenariusze testowe na podstawie przeanalizowanych danych.
        W rzeczywistości tutaj byłoby wywołanie modelu generatywnego.
        
        Args:
            analyzed_data: Przeanalizowane dane
            
        Returns:
            Lista scenariuszy testowych
        """
        test_scenarios = []
        insights = analyzed_data.get('combined_insights', [])
        
        # Generuj scenariusze dla każdego insightu
        for idx, insight in enumerate(insights, 1):
            insight_type = insight.get('type', 'requirement')
            description = insight.get('description', '')
            
            # Określ priorytet na podstawie typu
            if insight_type == 'requirement':
                priority = 'High'
            elif insight_type == 'functionality':
                priority = 'Medium'
            else:
                priority = 'Low'
            
            # Wygeneruj nazwę scenariusza na podstawie opisu
            scenario_name = self._generate_scenario_name(description, insight_type)
            
            # Wygeneruj kroki akcji na podstawie opisu
            step_action = self._generate_step_action(description, insight_type)
            
            # Wygeneruj oczekiwany rezultat
            expected_result = self._generate_expected_result(description, insight_type)
            
            test_scenarios.append({
                'test_case_id': f'TC_{idx:04d}',
                'scenario_name': scenario_name,
                'step_action': step_action,
                'requirement': description[:200],  # Ogranicz długość
                'expected_result': expected_result,
                'priority': priority,
                'status': 'Draft'
            })
        
        # Jeśli nie ma żadnych wniosków, użyj sekcji tekstu jako podstawy
        if not test_scenarios:
            text_analysis = analyzed_data.get('text_analysis', [])
            for idx, text_item in enumerate(text_analysis[:10], 1):  # Maksymalnie 10 scenariuszy
                section = text_item.get('section', f'Sekcja {idx}')
                content = text_item.get('content', '')
                
                # Użyj pierwszych 100 znaków jako wymaganie
                requirement = content[:100] + ('...' if len(content) > 100 else '')
                
                test_scenarios.append({
                    'test_case_id': f'TC_{idx:04d}',
                    'scenario_name': f'Test: {section}',
                    'step_action': 'Zweryfikuj funkcjonalność zgodnie z dokumentacją',
                    'requirement': requirement,
                    'expected_result': 'Funkcjonalność działa zgodnie z dokumentacją',
                    'priority': 'Medium',
                    'status': 'Draft'
                })
        
        # Jeśli nadal nie ma scenariuszy, stwórz jeden przykładowy
        if not test_scenarios:
            test_scenarios.append({
                'test_case_id': 'TC_0001',
                'scenario_name': 'Przykładowy scenariusz testowy',
                'step_action': 'Wykonaj podstawową akcję zgodnie z dokumentacją',
                'requirement': 'Wymaganie wyekstrahowane z dokumentacji',
                'expected_result': 'System działa poprawnie zgodnie z dokumentacją',
                'priority': 'High',
                'status': 'Draft'
            })
        
        return test_scenarios
    
    def _generate_scenario_name(self, description: str, insight_type: str) -> str:
        """Generuje nazwę scenariusza na podstawie opisu."""
        # Wyciągnij kluczowe słowa z opisu
        words = description.split()[:5]  # Pierwsze 5 słów
        name = ' '.join(words)
        
        # Skróć jeśli za długie
        if len(name) > 50:
            name = name[:47] + '...'
        
        if insight_type == 'requirement':
            return f'Test wymagania: {name}'
        elif insight_type == 'functionality':
            return f'Test funkcjonalności: {name}'
        elif insight_type == 'test_scenario':
            return f'Scenariusz: {name}'
        else:
            return f'Test: {name}'
    
    def _generate_step_action(self, description: str, insight_type: str) -> str:
        """Generuje krok akcji na podstawie opisu."""
        # Wyciągnij akcję z opisu
        action_keywords = ['wykonaj', 'otwórz', 'kliknij', 'wprowadź', 'wybierz', 'zapisz', 'usuń']
        
        for keyword in action_keywords:
            if keyword in description.lower():
                # Znajdź zdanie zawierające akcję
                sentences = re.split(r'[\.!?]\s+', description)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        return sentence.strip()[:150]  # Ogranicz długość
        
        # Domyślne akcje na podstawie typu
        if insight_type == 'requirement':
            return 'Zweryfikuj zgodność z wymaganiem'
        elif insight_type == 'functionality':
            return 'Przetestuj funkcjonalność'
        else:
            return 'Wykonaj akcję zgodnie z dokumentacją'
    
    def _generate_expected_result(self, description: str, insight_type: str) -> str:
        """Generuje oczekiwany rezultat na podstawie opisu."""
        # Szukaj słów wskazujących na rezultat
        result_keywords = ['powinien', 'powinna', 'powinno', 'musi', 'oczekiwany', 'rezultat']
        
        for keyword in result_keywords:
            if keyword in description.lower():
                # Znajdź zdanie zawierające rezultat
                sentences = re.split(r'[\.!?]\s+', description)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        return sentence.strip()[:150]
        
        # Domyślne rezultaty
        if insight_type == 'requirement':
            return 'Wymaganie jest spełnione'
        elif insight_type == 'functionality':
            return 'Funkcjonalność działa poprawnie'
        else:
            return 'Oczekiwany rezultat zgodny z dokumentacją'
    
    def save_results(self, test_scenarios: List[Dict], results_dir: Path, task_id: str) -> Path:
        """
        Zapisuje wyniki do pliku Excel.
        
        Args:
            test_scenarios: Lista scenariuszy testowych
            results_dir: Katalog wyników
            task_id: ID zadania
            
        Returns:
            Ścieżka do zapisanego pliku
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Scenariusze Testowe"
            
            # Nagłówki
            headers = [
                'Test Case ID',
                'Nazwa scenariusza',
                'Krok do wykonania',
                'Wymaganie',
                'Rezultat',
                'Priorytet',
                'Status'
            ]
            
            # Styl nagłówków
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Dane
            for row_idx, scenario in enumerate(test_scenarios, 2):
                ws.cell(row=row_idx, column=1, value=scenario.get('test_case_id', ''))
                ws.cell(row=row_idx, column=2, value=scenario.get('scenario_name', ''))
                ws.cell(row=row_idx, column=3, value=scenario.get('step_action', ''))
                ws.cell(row=row_idx, column=4, value=scenario.get('requirement', ''))
                ws.cell(row=row_idx, column=5, value=scenario.get('expected_result', ''))
                ws.cell(row=row_idx, column=6, value=scenario.get('priority', ''))
                ws.cell(row=row_idx, column=7, value=scenario.get('status', ''))
            
            # Dostosuj szerokość kolumn
            column_widths = [15, 30, 40, 20, 40, 12, 12]
            for col_idx, width in enumerate(column_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
            
            # Zapisz plik
            results_dir.mkdir(parents=True, exist_ok=True)
            result_file = results_dir / f"wyniki_{task_id}.xlsx"
            wb.save(str(result_file))
            
            return result_file
        
        except ImportError:
            # Fallback: zapisz jako JSON, jeśli openpyxl nie jest dostępne
            results_dir.mkdir(parents=True, exist_ok=True)
            result_file = results_dir / f"wyniki_{task_id}.json"
            
            with open(result_file, 'w', encoding='utf-8') as f:
                json.dump(test_scenarios, f, ensure_ascii=False, indent=2)
            
            return result_file
    
    def _load_prompt(self, prompt_file: str) -> str:
        """Wczytuje prompt z pliku."""
        prompt_path = Path(prompt_file)
        if not prompt_path.exists():
            raise FileNotFoundError(f"Plik promptu nie istnieje: {prompt_file}")
        with open(prompt_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _load_settings(self) -> Dict:
        """Wczytuje ustawienia z pliku settings.txt."""
        settings = {
            'temperature': 0.2,
            'top_p': 0.9,
            'top_k': 40,
            'max_tokens': 2048
        }
        settings_path = Path('settings.txt')
        if settings_path.exists():
            with open(settings_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        if '=' in line:
                            key, value = line.split('=', 1)
                            key = key.strip()
                            value = value.strip()
                            try:
                                if '.' in value:
                                    settings[key] = float(value)
                                else:
                                    settings[key] = int(value)
                            except ValueError:
                                pass
        return settings
    
    def _call_ollama(self, prompt: str, system_prompt: str = None, max_retries: int = 3) -> str:
        """Wywołuje Ollama API z promptem."""
        import time
        settings = self._load_settings()
        api_url = f"{self.ollama_url}/api/generate"
        
        payload = {
            "model": self.ollama_model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": settings.get('temperature', 0.2),
                "top_p": settings.get('top_p', 0.9),
                "top_k": settings.get('top_k', 40),
                "num_predict": settings.get('max_tokens', 2048)
            }
        }
        
        if system_prompt:
            payload["system"] = system_prompt
        
        for attempt in range(max_retries):
            try:
                response = requests.post(api_url, json=payload, timeout=300)
                if response.status_code == 200:
                    result = response.json()
                    return result.get('response', '').strip()
                elif response.status_code == 500:
                    error_text = response.text.lower()
                    # Wykryj błędy pamięci GPU
                    if 'resource' in error_text or 'memory' in error_text or 'stopped' in error_text:
                        print(f"  ⚠️ Błąd pamięci GPU (próba {attempt + 1}/{max_retries})")
                        print(f"  💡 Model mógł zostać zatrzymany z powodu braku pamięci VRAM")
                        if attempt < max_retries - 1:
                            print(f"  🔄 Czekam 10s przed ponowną próbą...")
                            time.sleep(10)  # Dłuższe oczekiwanie przy błędach pamięci
                            continue
                    # Wykryj błędy kontekstu
                    elif 'context' in error_text or 'token' in error_text or 'length' in error_text:
                        print(f"  ⚠️ Błąd kontekstu/tokenów (próba {attempt + 1}/{max_retries})")
                        # Zwróć specjalny błąd który można obsłużyć wyżej
                        raise ContextLengthError(f"Przekroczono limit kontekstu: {response.text[:200]}")
                    else:
                        print(f"Błąd serwera Ollama (500, próba {attempt + 1}/{max_retries}): {response.text[:300]}")
                else:
                    print(f"Błąd Ollama (próba {attempt + 1}/{max_retries}): {response.status_code} - {response.text[:300]}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
            except ContextLengthError:
                # Przepuść błąd kontekstu do obsługi wyżej
                raise
            except requests.exceptions.Timeout:
                print(f"  ⚠️ Timeout Ollama (próba {attempt + 1}/{max_retries}) - model może być przeciążony")
                if attempt < max_retries - 1:
                    time.sleep(5)
            except requests.exceptions.RequestException as e:
                print(f"Błąd połączenia z Ollama (próba {attempt + 1}/{max_retries}): {e}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
        
        raise Exception("Nie udało się uzyskać odpowiedzi z Ollama po wszystkich próbach")
    
    def _extract_sections_from_content(self, extracted_data: Dict) -> Dict[str, Dict]:
        """
        Ekstrahuje zawartość dokumentu z numeracją sekcji i metadanymi obrazów.
        Używa rzeczywistych sekcji z dokumentu (nagłówki lub automatyczna numeracja).
        
        Returns:
            Słownik: nazwa_sekcji -> {'content': zawartość_sekcji, 'images': lista_obrazów, 'section_index': int}
        """
        sections = {}
        image_descriptions = extracted_data.get('image_descriptions', {})
        
        for section_index, text_item in enumerate(extracted_data.get('text', []), 1):
            section_title = text_item.get('section', f'Sekcja {section_index}')
            content = text_item.get('content', '')
            
            # Znajdź obrazy w tej sekcji
            image_placeholders = text_item.get('image_placeholders', [])
            section_images = []
            for img_placeholder in image_placeholders:
                img_filename = img_placeholder.get('filename', '')
                if img_filename:
                    img_info = {
                        'filename': img_filename,
                        'description': image_descriptions.get(img_filename, 'Brak opisu'),
                        'section': section_title
                    }
                    section_images.append(img_info)
            
            # Zapisz sekcję
            sections[section_title] = {
                'content': content,
                'images': section_images,
                'section_index': section_index
            }
        
        return sections
    
    def _get_document_fragments(self, sections: Dict[str, Dict], section_names: List[str]) -> str:
        """
        Wyciąga fragmenty dokumentacji dla określonych sekcji wraz z metadanymi obrazów.
        
        Args:
            sections: Słownik wszystkich sekcji (z metadanymi)
            section_names: Lista nazw sekcji do wyciągnięcia
            
        Returns:
            Połączona zawartość wybranych sekcji z opisami obrazów
        """
        fragments = []
        # Sortuj sekcje według ich indeksu w dokumencie
        sorted_sections = sorted(
            [(name, sections[name]) for name in section_names if name in sections],
            key=lambda x: x[1].get('section_index', 999)
        )
        
        for section_name, section_data in sorted_sections:
            content = section_data.get('content', '')
            images = section_data.get('images', [])
            
            # Dodaj zawartość sekcji
            fragment = f"=== SEKCJA: {section_name} ===\n{content}"
            
            # Dodaj metadane obrazów jeśli są
            if images:
                fragment += "\n\n[OBRAZY W TEJ SEKCJI:]\n"
                for img in images:
                    fragment += f"- {img.get('filename', '')}: {img.get('description', 'Brak opisu')}\n"
            
            fragments.append(fragment)
        return '\n\n'.join(fragments)
    
    def _split_documentation_into_chunks(self, doc_text: str, max_tokens: int = 12000) -> List[str]:
        """
        Dzieli długą dokumentację na mniejsze fragmenty (chunki) mieszczące się w limicie tokenów.
        
        Args:
            doc_text: Pełny tekst dokumentacji
            max_tokens: Maksymalna liczba tokenów na chunk (domyślnie 12000 dla limitu 16k)
            
        Returns:
            Lista chunków dokumentacji
        """
        # Przybliżone oszacowanie: 1 token ≈ 4 znaki dla języka polskiego
        chars_per_token = 4
        max_chars = max_tokens * chars_per_token
        
        # Jeśli dokumentacja mieści się w limicie, zwróć jako jeden chunk
        if len(doc_text) <= max_chars:
            return [doc_text]
        
        # Podziel dokument po sekcjach
        sections = doc_text.split('## ')
        chunks = []
        current_chunk = ""
        
        for i, section in enumerate(sections):
            # Dodaj z powrotem separator dla nie-pierwszej sekcji
            section_text = section if i == 0 else f"## {section}"
            
            # Jeśli pojedyncza sekcja jest większa niż limit, dziel ją na akapity
            if len(section_text) > max_chars:
                # Zapisz poprzedni chunk jeśli nie jest pusty
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                
                # Dziel dużą sekcję na akapity
                paragraphs = section_text.split('\n\n')
                for para in paragraphs:
                    if len(current_chunk) + len(para) + 2 > max_chars:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = para + "\n\n"
                    else:
                        current_chunk += para + "\n\n"
            else:
                # Sprawdź czy dodanie sekcji przekroczy limit
                if len(current_chunk) + len(section_text) > max_chars:
                    # Zapisz obecny chunk i zacznij nowy
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = section_text + "\n\n"
                else:
                    current_chunk += section_text + "\n\n"
        
        # Dodaj ostatni chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [doc_text]
    
    def _format_source_sections(self, section_names: List[str]) -> str:
        """
        Formatuje listę nazw sekcji jako czytelny opis.
        Np. ["Wstęp", "Instalacja", "Konfiguracja"] -> "sekcje: Wstęp, Instalacja, Konfiguracja"
        
        Args:
            section_names: Lista nazw sekcji
            
        Returns:
            Sformatowany opis sekcji
        """
        if not section_names:
            return "cała dokumentacja"
        
        # Usuń duplikaty zachowując kolejność
        unique_sections = []
        seen = set()
        for section in section_names:
            if section and section not in seen:
                unique_sections.append(section)
                seen.add(section)
        
        if len(unique_sections) == 1:
            return f"sekcja: {unique_sections[0]}"
        
        # Ogranicz do 5 sekcji, jeśli więcej dodaj "..."
        if len(unique_sections) > 5:
            return f"sekcje: {', '.join(unique_sections[:5])}, ... (łącznie {len(unique_sections)} sekcji)"
        
        return f"sekcje: {', '.join(unique_sections)}"

    def _normalize_steps(self, steps: List[Any]) -> List[Dict]:
        """
        Normalizuje strukturę kroków zwracaną przez modele LLM.
        Zapewnia, że każdy krok jest słownikiem z kluczami:
        step_number, action, expected_result.
        """
        normalized_steps: List[Dict] = []
        for idx, step in enumerate(steps, 1):
            if isinstance(step, dict):
                normalized_step = dict(step)
                normalized_step.setdefault('step_number', normalized_step.get('order', idx) or idx)
                action = normalized_step.get('action') or normalized_step.get('description') or normalized_step.get('name')
                normalized_step['action'] = action or f'Krok {idx} - brak opisu'
                expected = normalized_step.get('expected_result') or normalized_step.get('result')
                normalized_step['expected_result'] = expected or 'Brak oczekiwanego rezultatu - weryfikacja ręczna'
                normalized_steps.append(normalized_step)
            elif isinstance(step, str):
                normalized_steps.append({
                    'step_number': idx,
                    'action': step.strip() or f'Krok {idx} - brak treści',
                    'expected_result': 'Weryfikacja ręczna zgodnie z dokumentacją'
                })
            else:
                # Nieobsługiwany typ kroku – zapisz jako informację tekstową
                normalized_steps.append({
                    'step_number': idx,
                    'action': f'Krok {idx}: {repr(step)}',
                    'expected_result': 'Weryfikacja ręczna'
                })
        return normalized_steps
    
    def stage1_generate_test_paths(self, extracted_data: Dict, processing_dir: Path, results_dir: Path = None, task_id: str = None) -> List[Dict]:
        """
        ETAP 1: Generuje ścieżki testowe na podstawie dokumentacji.
        Dla długich dokumentów dzieli na chunki i przetwarza osobno.
        
        Args:
            extracted_data: Ekstrahowane dane z dokumentu
            processing_dir: Katalog przetwarzania
            
        Returns:
            Lista ścieżek testowych w formacie JSON
        """
        import time
        
        # Przygotuj pełną dokumentację
        full_documentation = []
        for text_item in extracted_data.get('text', []):
            section = text_item.get('section', '')
            content = text_item.get('content', '')
            if section:
                full_documentation.append(f"## {section}\n{content}")
            else:
                full_documentation.append(content)
        
        doc_text = '\n\n'.join(full_documentation)
        
        # Wczytaj prompt dla etapu 1
        prompt_template = self._load_prompt('prompt1.txt')
        
        # Podziel dokumentację na chunki (limit 12k tokenów = ~48k znaków)
        doc_chunks = self._split_documentation_into_chunks(doc_text, max_tokens=12000)
        
        # Aktualizuj statystyki przetwarzania
        self.processing_stats['current_stage'] = 1
        self.processing_stats['total_chunks'] = len(doc_chunks)
        self.processing_stats['processed_chunks'] = 0
        self.processing_stats['start_time'] = time.time()
        
        print(f"ETAP 1: Generowanie ścieżek testowych... (Dokumentacja podzielona na {len(doc_chunks)} fragmentów)")
        
        all_paths = []
        path_id_counter = 1
        
        # Przetwarzaj każdy chunk osobno
        for chunk_idx, chunk in enumerate(doc_chunks, 1):
            chunk_start_time = time.time()
            
            # Oblicz dynamiczny ETA
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Przetwarzanie fragmentu {chunk_idx}/{len(doc_chunks)}...{eta_str}")
            
            # Elastyczna liczba ścieżek - model sam decyduje ile potrzeba
            full_prompt = f"{prompt_template}\n\nDOKUMENTACJA (Fragment {chunk_idx}/{len(doc_chunks)}):\n{chunk}\n\nPAMIĘTAJ: Zwracasz TYLKO JSON (bez żadnego tekstu przed lub po). Wygeneruj ścieżki testowe pokrywające WSZYSTKIE funkcjonalności z tego fragmentu dokumentacji. Liczba ścieżek zależy od zawartości - może być 5 lub 50, ważne jest pełne pokrycie."
            
            # Wywołaj Ollama
            response = self._call_ollama(full_prompt)
            
            # Zapisz czas przetwarzania chunka
            chunk_time = time.time() - chunk_start_time
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = chunk_idx
            
            if not response or len(response.strip()) < 10:
                print(f"  OSTRZEŻENIE: Pusty response dla fragmentu {chunk_idx}, pomijam...")
                continue
            
            # Parsuj JSON z odpowiedzi
            try:
                # Wyciągnij JSON z odpowiedzi (może być otoczony tekstem)
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                else:
                    # Spróbuj znaleźć JSON zaczynający się od [
                    json_start = response.find('[')
                    if json_start != -1:
                        # Znajdź zbalansowane nawiasy kwadratowe
                        bracket_count = 0
                        json_end = json_start
                        for i in range(json_start, len(response)):
                            if response[i] == '[':
                                bracket_count += 1
                            elif response[i] == ']':
                                bracket_count -= 1
                                if bracket_count == 0:
                                    json_end = i + 1
                                    break
                        json_str = response[json_start:json_end]
                    else:
                        print(f"  OSTRZEŻENIE: Nie znaleziono JSON w odpowiedzi dla fragmentu {chunk_idx}")
                        continue
                
                paths = json.loads(json_str)
                
                # Debug: pokaż co zwrócił model
                if paths and len(paths) > 0:
                    print(f"  DEBUG: Typ pierwszego elementu: {type(paths[0])}, wartość: {str(paths[0])[:200]}")
                
                # Upewnij się, że paths jest listą słowników
                if not isinstance(paths, list):
                    paths = [paths] if isinstance(paths, dict) else []
                
                # Jeśli model zwrócił listę stringów, spróbuj je przekonwertować na słowniki
                if paths and isinstance(paths[0], str):
                    print(f"  INFO: Model zwrócił listę stringów, konwertuję na słowniki...")
                    converted_paths = []
                    for idx, path_str in enumerate(paths):
                        converted_paths.append({
                            'title': path_str,
                            'description': path_str,
                            'type': 'happy_path',
                            'source_sections': [],
                            'border_conditions': []
                        })
                    paths = converted_paths
                
                # Sprawdź czy każdy element jest słownikiem i ma wymagane pola
                for i, path in enumerate(paths):
                    if isinstance(path, dict):
                        # Spróbuj znaleźć tytuł w różnych polach (model może używać różnych nazw)
                        title = (path.get('title') or path.get('name') or path.get('test_name') or 
                                 path.get('nazwa') or path.get('description') or path.get('opis') or
                                 path.get('scenario_name') or path.get('test_case'))
                        
                        if title:
                            # Ustaw tytuł jeśli go nie było
                            if 'title' not in path:
                                path['title'] = title
                            
                            # Upewnij się, że ma unikalne ID (nadpisujemy zawsze)
                            path['id'] = f"PATH_{path_id_counter:03d}"
                            path_id_counter += 1
                            
                            # Upewnij się, że source_sections istnieje
                            if 'source_sections' not in path:
                                if 'source_pages' in path:
                                    path['source_sections'] = []
                                else:
                                    path['source_sections'] = []
                            if 'type' not in path:
                                path['type'] = 'happy_path'
                            if 'description' not in path:
                                path['description'] = path.get('title', '')
                            if 'border_conditions' not in path:
                                path['border_conditions'] = []
                            all_paths.append(path)
                        else:
                            print(f"  Ostrzeżenie: Ścieżka {i} w fragmencie {chunk_idx} nie ma tytułu (title/name/description), pomijam: {list(path.keys())}")
                    else:
                        print(f"  Ostrzeżenie: Ścieżka {i} w fragmencie {chunk_idx} nie jest słownikiem, pomijam")
                
                print(f"  Fragment {chunk_idx}: Wygenerowano {len(paths)} ścieżek")
                
            except json.JSONDecodeError as e:
                print(f"  Błąd parsowania JSON fragmentu {chunk_idx}: {e}")
                print(f"  Odpowiedź Ollama: {response[:300]}")
                continue
        
        # Sprawdź czy udało się wygenerować jakiekolwiek ścieżki
        if len(all_paths) == 0:
            raise Exception("Nie udało się wygenerować żadnych ścieżek testowych z żadnego fragmentu")
        
        # Zapisz wszystkie ścieżki do pliku tymczasowego
        paths_file = processing_dir / "sciezki_testowe.txt"
        with open(paths_file, 'w', encoding='utf-8') as f:
            json.dump(all_paths, f, ensure_ascii=False, indent=2)
        
        # Zapisz również do results_dir jako artefakt do pobrania
        if results_dir and task_id:
            results_dir.mkdir(parents=True, exist_ok=True)
            artifact_file = results_dir / f"etap1_sciezki_testowe_{task_id}.json"
            with open(artifact_file, 'w', encoding='utf-8') as f:
                json.dump(all_paths, f, ensure_ascii=False, indent=2)
            print(f"  Artefakt Etapu 1 zapisany: {artifact_file.name}")
        
        print(f"ETAP 1: ŁĄCZNIE wygenerowano {len(all_paths)} ścieżek testowych z {len(doc_chunks)} fragmentów")
        return all_paths
    
    def stage2_generate_scenarios(self, extracted_data: Dict, test_paths: List[Dict], processing_dir: Path, results_dir: Path = None, task_id: str = None) -> List[Dict]:
        """
        ETAP 2: Generuje scenariusze testowe z walidacjami.
        Dla długich dokumentów dzieli na chunki.
        
        Args:
            extracted_data: Ekstrahowane dane z dokumentu
            test_paths: Lista ścieżek testowych z etapu 1
            processing_dir: Katalog przetwarzania
            
        Returns:
            Lista scenariuszy testowych
        """
        import time
        
        # Przygotuj pełną dokumentację
        full_documentation = []
        for text_item in extracted_data.get('text', []):
            section = text_item.get('section', '')
            content = text_item.get('content', '')
            if section:
                full_documentation.append(f"## {section}\n{content}")
            else:
                full_documentation.append(content)
        
        doc_text = '\n\n'.join(full_documentation)
        
        # Wczytaj prompt dla etapu 2
        prompt_template = self._load_prompt('prompt2.txt')
        
        # Podziel dokumentację na chunki (limit 12k tokenów)
        doc_chunks = self._split_documentation_into_chunks(doc_text, max_tokens=12000)
        
        # Aktualizuj statystyki przetwarzania
        self.processing_stats['current_stage'] = 2
        self.processing_stats['total_chunks'] = len(doc_chunks)
        self.processing_stats['processed_chunks'] = 0
        
        print(f"ETAP 2: Generowanie scenariuszy testowych... (Dokumentacja podzielona na {len(doc_chunks)} fragmentów)")
        
        all_scenarios = []
        scenario_id_counter = 1
        
        # Przygotuj skróconą listę ścieżek (tylko tytuły i ID) aby zmniejszyć prompt
        paths_summary = [{"id": p.get("id", ""), "title": p.get("title", ""), "type": p.get("type", "")} for p in test_paths]
        paths_json = json.dumps(paths_summary, ensure_ascii=False, indent=2)
        
        # Przetwarzaj każdy chunk osobno
        for chunk_idx, chunk in enumerate(doc_chunks, 1):
            chunk_start_time = time.time()
            
            # Oblicz dynamiczny ETA
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Przetwarzanie fragmentu {chunk_idx}/{len(doc_chunks)}...{eta_str}")
            
            # Elastyczna liczba scenariuszy - model sam decyduje
            full_prompt = f"{prompt_template}\n\nDOKUMENTACJA (Fragment {chunk_idx}/{len(doc_chunks)}):\n{chunk}\n\nŚCIEŻKI TESTOWE (skrót):\n{paths_json}\n\nPAMIĘTAJ: Zwracasz TYLKO JSON (bez żadnego tekstu przed lub po). Wygeneruj scenariusze testowe pokrywające funkcjonalności z tego fragmentu. Dla każdej funkcjonalności uwzględnij: happy path, przypadki negatywne (błędne dane), walidacje pól wymaganych. Liczba scenariuszy zależy od zawartości dokumentacji."
            
            # Wywołaj Ollama
            response = self._call_ollama(full_prompt)
            
            # Zapisz czas przetwarzania chunka
            chunk_time = time.time() - chunk_start_time
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = chunk_idx
            
            if not response or len(response.strip()) < 10:
                print(f"  OSTRZEŻENIE: Pusty response dla fragmentu {chunk_idx}, pomijam...")
                continue
            
            # Parsuj JSON
            try:
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                else:
                    # Spróbuj znaleźć JSON zaczynający się od [
                    json_start = response.find('[')
                    if json_start != -1:
                        bracket_count = 0
                        json_end = json_start
                        for i in range(json_start, len(response)):
                            if response[i] == '[':
                                bracket_count += 1
                            elif response[i] == ']':
                                bracket_count -= 1
                                if bracket_count == 0:
                                    json_end = i + 1
                                    break
                        json_str = response[json_start:json_end]
                    else:
                        print(f"  OSTRZEŻENIE: Nie znaleziono JSON w odpowiedzi dla fragmentu {chunk_idx}")
                        continue
                
                scenarios = json.loads(json_str)
                
                # Upewnij się, że scenarios jest listą słowników
                if not isinstance(scenarios, list):
                    scenarios = [scenarios] if isinstance(scenarios, dict) else []
                
                # Sprawdź czy każdy element jest słownikiem i ma wymagane pola
                for i, scenario in enumerate(scenarios):
                    if isinstance(scenario, dict):
                        # Spróbuj znaleźć tytuł w różnych polach (model może używać różnych nazw)
                        title = scenario.get('title') or scenario.get('name') or scenario.get('nazwa') or scenario.get('scenario_name') or scenario.get('description') or scenario.get('opis')
                        
                        if title:
                            # Ustaw tytuł jeśli go nie było
                            if 'title' not in scenario:
                                scenario['title'] = title
                            
                            # Upewnij się, że ma unikalne ID (nadpisujemy zawsze)
                            scenario['scenario_id'] = f"SCEN_{scenario_id_counter:03d}"
                            scenario_id_counter += 1
                            
                            # Upewnij się, że source_sections istnieje
                            if 'source_sections' not in scenario:
                                if 'source_pages' in scenario:
                                    scenario['source_sections'] = []
                                else:
                                    scenario['source_sections'] = []
                            if 'priority' not in scenario:
                                scenario['priority'] = 'Medium'
                            if 'type' not in scenario:
                                scenario['type'] = 'positive'
                            all_scenarios.append(scenario)
                        else:
                            print(f"  Ostrzeżenie: Scenariusz {i} w fragmencie {chunk_idx} nie ma tytułu (title/name/scenario_name/description), pomijam: {list(scenario.keys())}")
                    else:
                        print(f"  Ostrzeżenie: Scenariusz {i} w fragmencie {chunk_idx} nie jest słownikiem, pomijam")
                
                print(f"  Fragment {chunk_idx}: Wygenerowano {len(scenarios)} scenariuszy")
                
            except json.JSONDecodeError as e:
                print(f"  Błąd parsowania JSON fragmentu {chunk_idx}: {e}")
                print(f"  Odpowiedź Ollama: {response[:300]}")
                continue
        
        # Sprawdź czy udało się wygenerować jakiekolwiek scenariusze
        if len(all_scenarios) == 0:
            raise Exception("Nie udało się wygenerować żadnych scenariuszy testowych z żadnego fragmentu")
        
        if len(all_scenarios) < 30:
            print(f"  UWAGA: Wygenerowano tylko {len(all_scenarios)} scenariuszy")
        
        # Zapisz scenariusze do pliku tymczasowego
        scenarios_file = processing_dir / "scenariusze_testowe.txt"
        with open(scenarios_file, 'w', encoding='utf-8') as f:
            json.dump(all_scenarios, f, ensure_ascii=False, indent=2)
        
        # Zapisz również do results_dir jako artefakt do pobrania
        if results_dir and task_id:
            results_dir.mkdir(parents=True, exist_ok=True)
            artifact_file = results_dir / f"etap2_scenariusze_{task_id}.json"
            with open(artifact_file, 'w', encoding='utf-8') as f:
                json.dump(all_scenarios, f, ensure_ascii=False, indent=2)
            print(f"  Artefakt Etapu 2 zapisany: {artifact_file.name}")
        
        print(f"ETAP 2: ŁĄCZNIE wygenerowano {len(all_scenarios)} scenariuszy testowych z {len(doc_chunks)} fragmentów")
        return all_scenarios
    
    def _process_batch_with_fallback(self, batch: List[Dict], prompt_template: str, full_doc: str, 
                                       batch_idx: int, total_batches: int, current_batch_size: int = None) -> List[Dict]:
        """
        Przetwarza batch scenariuszy z fallbackiem na mniejsze batche przy błędach kontekstu.
        
        Args:
            batch: Lista scenariuszy do przetworzenia
            prompt_template: Szablon promptu
            full_doc: Skrócona dokumentacja
            batch_idx: Indeks batcha
            total_batches: Całkowita liczba batchy
            current_batch_size: Aktualny rozmiar batcha (do fallbacku)
            
        Returns:
            Lista przetworzonych scenariuszy lub pusta lista przy błędzie
        """
        if current_batch_size is None:
            current_batch_size = len(batch)
        
        # Skróć dokumentację jeśli batch jest duży
        doc_limit = 32000 if current_batch_size <= 3 else 24000 if current_batch_size <= 5 else 16000
        truncated_doc = full_doc[:doc_limit] if len(full_doc) > doc_limit else full_doc
        if len(full_doc) > doc_limit:
            truncated_doc += "\n\n[...dokumentacja skrócona ze względu na limit kontekstu...]"
        
        batch_scenarios_json = json.dumps(batch, ensure_ascii=False, indent=2)
        
        full_prompt = f"""{prompt_template}

DOKUMENTACJA:
{truncated_doc}

SCENARIUSZE DO PRZETWORZENIA (batch {batch_idx}/{total_batches}):
{batch_scenarios_json}

PAMIĘTAJ: Zwracasz TYLKO tablicę JSON z {len(batch)} scenariuszami (bez żadnego tekstu przed lub po). 
Każdy scenariusz MUSI zawierać: scenario_id, test_case_id, scenario_name, steps (lista kroków).
Każdy krok MUSI mieć: step_number, action, expected_result.
Liczba kroków dostosowana do złożoności scenariusza (minimum 3)."""
        
        try:
            response = self._call_ollama(full_prompt)
            
            if not response or len(response.strip()) < 10:
                raise Exception("Ollama zwróciła pustą odpowiedź")
            
            # Parsuj tablicę JSON
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
            else:
                json_start = response.find('[')
                if json_start != -1:
                    bracket_count = 0
                    json_end = json_start
                    for i in range(json_start, len(response)):
                        if response[i] == '[':
                            bracket_count += 1
                        elif response[i] == ']':
                            bracket_count -= 1
                            if bracket_count == 0:
                                json_end = i + 1
                                break
                    json_str = response[json_start:json_end]
                else:
                    json_str = response
            
            batch_results = json.loads(json_str)
            
            if not isinstance(batch_results, list):
                batch_results = [batch_results] if isinstance(batch_results, dict) else []
            
            return batch_results
            
        except ContextLengthError as e:
            # Błąd kontekstu - podziel batch na mniejsze
            print(f"    ⚠️ Błąd kontekstu dla batcha {len(batch)} scenariuszy - dzielę na mniejsze...")
            
            if len(batch) == 1:
                # Nie można już podzielić - spróbuj z minimalną dokumentacją
                print(f"    ⚠️ Próbuję z minimalną dokumentacją dla pojedynczego scenariusza...")
                minimal_doc = full_doc[:8000] + "\n\n[...dokumentacja mocno skrócona...]"
                try:
                    return self._process_batch_with_fallback(
                        batch, prompt_template, minimal_doc, batch_idx, total_batches, 1
                    )
                except Exception:
                    print(f"    ❌ Nie udało się przetworzyć scenariusza - oznaczam jako błąd")
                    return []
            
            # Podziel batch na dwie połowy
            mid = len(batch) // 2
            first_half = batch[:mid]
            second_half = batch[mid:]
            
            results = []
            if first_half:
                results.extend(self._process_batch_with_fallback(
                    first_half, prompt_template, full_doc, batch_idx, total_batches, len(first_half)
                ))
            if second_half:
                results.extend(self._process_batch_with_fallback(
                    second_half, prompt_template, full_doc, batch_idx, total_batches, len(second_half)
                ))
            
            return results
            
        except Exception as e:
            print(f"    ❌ Błąd przetwarzania batcha: {e}")
            
            # Przy innych błędach też spróbuj podzielić batch
            if len(batch) > 1:
                print(f"    🔄 Próbuję podzielić batch {len(batch)} scenariuszy...")
                mid = len(batch) // 2
                results = []
                try:
                    results.extend(self._process_batch_with_fallback(
                        batch[:mid], prompt_template, full_doc, batch_idx, total_batches, mid
                    ))
                    results.extend(self._process_batch_with_fallback(
                        batch[mid:], prompt_template, full_doc, batch_idx, total_batches, len(batch) - mid
                    ))
                    return results
                except Exception:
                    pass
            
            return []
    
    def stage3_generate_detailed_steps(self, extracted_data: Dict, scenarios: List[Dict], processing_dir: Path, results_dir: Path, task_id: str) -> Path:
        """
        ETAP 3: Generuje szczegółowe kroki testowe z BATCH PROCESSING.
        Przetwarza kilka scenariuszy naraz zamiast osobno.
        
        Args:
            extracted_data: Ekstrahowane dane z dokumentu
            scenarios: Lista scenariuszy z etapu 2
            processing_dir: Katalog przetwarzania
            results_dir: Katalog wyników
            task_id: ID zadania
            
        Returns:
            Ścieżka do pliku Excel z wynikami
        """
        import time
        
        # Przygotuj mapę sekcji
        sections = self._extract_sections_from_content(extracted_data)
        
        # Wczytaj prompt dla etapu 3
        prompt_template = self._load_prompt('prompt3.txt')
        
        # Przygotuj listę wszystkich szczegółowych scenariuszy
        all_detailed_scenarios = []
        
        # BATCH PROCESSING: Przetwarzaj scenariusze w grupach po 5
        BATCH_SIZE = 5
        batches = [scenarios[i:i + BATCH_SIZE] for i in range(0, len(scenarios), BATCH_SIZE)]
        
        # Aktualizuj statystyki przetwarzania
        self.processing_stats['current_stage'] = 3
        self.processing_stats['total_chunks'] = len(batches)
        self.processing_stats['processed_chunks'] = 0
        
        print(f"ETAP 3: Generowanie szczegółowych kroków dla {len(scenarios)} scenariuszy w {len(batches)} batchach...")
        
        # Przygotuj skróconą dokumentację (dla wszystkich scenariuszy)
        full_doc = '\n\n'.join([f"## {item.get('section', '')}\n{item.get('content', '')}" 
                               for item in extracted_data.get('text', [])])
        # Ogranicz rozmiar dokumentacji do ~8k tokenów
        max_doc_chars = 32000
        if len(full_doc) > max_doc_chars:
            full_doc = full_doc[:max_doc_chars] + "\n\n[...dokumentacja skrócona...]"
        
        # Przetwarzaj każdy batch z obsługą błędów kontekstu
        for batch_idx, batch in enumerate(batches, 1):
            batch_start_time = time.time()
            
            # Oblicz dynamiczny ETA
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Przetwarzanie batcha {batch_idx}/{len(batches)} ({len(batch)} scenariuszy)...{eta_str}")
            
            # Przetwórz batch z fallbackiem na mniejsze batche
            batch_results = self._process_batch_with_fallback(
                batch, prompt_template, full_doc, batch_idx, len(batches)
            )
            
            # Zapisz czas przetwarzania batcha
            batch_time = time.time() - batch_start_time
            self.processing_stats['chunk_times'].append(batch_time)
            self.processing_stats['processed_chunks'] = batch_idx
            
            if batch_results:
                print(f"    Otrzymano {len(batch_results)} scenariuszy z batcha")
                
                # Przetwórz każdy scenariusz z batcha
                for scen_idx, detailed_scenario in enumerate(batch_results):
                    if not isinstance(detailed_scenario, dict):
                        continue
                    
                    # Pobierz oryginalny scenariusz z batcha
                    orig_scenario = batch[scen_idx] if scen_idx < len(batch) else {}
                    source_sections = orig_scenario.get('source_sections', [])
                    
                    # Upewnij się, że source_sections są zapisane
                    if 'source_sections' not in detailed_scenario:
                        detailed_scenario['source_sections'] = source_sections
                    
                    # Upewnij się, że są kroki
                    steps = detailed_scenario.get('steps', [])
                    if not isinstance(steps, list):
                        steps = []
                    steps = self._normalize_steps(steps)
                    detailed_scenario['steps'] = steps
                    
                    # Sprawdź czy jest co najmniej 3 kroki
                    if len(detailed_scenario['steps']) < 3:
                        while len(detailed_scenario['steps']) < 3:
                            step_num = len(detailed_scenario['steps']) + 1
                            detailed_scenario['steps'].append({
                                'step_number': step_num,
                                'action': f'Krok {step_num} - wymagana ręczna weryfikacja',
                                'expected_result': 'Wymagana ręczna weryfikacja zgodnie z dokumentacją'
                            })
                        detailed_scenario['steps'] = self._normalize_steps(detailed_scenario['steps'])
                    
                    # Upewnij się, że test_case_id istnieje
                    if 'test_case_id' not in detailed_scenario:
                        detailed_scenario['test_case_id'] = f'TC_{len(all_detailed_scenarios) + 1:04d}'
                    
                    # Upewnij się, że scenario_name istnieje
                    if 'scenario_name' not in detailed_scenario:
                        detailed_scenario['scenario_name'] = orig_scenario.get('title', orig_scenario.get('scenario_name', f'Scenariusz'))
                    
                    all_detailed_scenarios.append(detailed_scenario)
            else:
                # Fallback: dodaj błędne scenariusze
                print(f"  ⚠️ Batch {batch_idx} nie zwrócił wyników - dodaję scenariusze z błędem")
                for scenario in batch:
                    scenario_id = scenario.get('scenario_id', f'SCEN_{len(all_detailed_scenarios) + 1:03d}')
                    all_detailed_scenarios.append({
                        'scenario_id': scenario_id,
                        'test_case_id': f'TC_{len(all_detailed_scenarios) + 1:04d}',
                        'scenario_name': scenario.get('title', 'Błąd generowania'),
                        'source_sections': scenario.get('source_sections', []),
                        'priority': scenario.get('priority', 'Medium'),
                        'status': 'Error',
                        'steps': [{
                            'step_number': 1,
                            'action': 'Scenariusz wymaga ręcznego uzupełnienia',
                            'expected_result': 'Wymagana ręczna weryfikacja'
                        }]
                    })
        
        # Weryfikacja: sprawdź czy wszystkie scenariusze zostały przetworzone
        processed_scenario_ids = {s.get('scenario_id') for s in all_detailed_scenarios}
        original_scenario_ids = {s.get('scenario_id') for s in scenarios}
        missing_scenarios = original_scenario_ids - processed_scenario_ids
        
        if missing_scenarios:
            print(f"  UWAGA: Nie przetworzono {len(missing_scenarios)} scenariuszy: {missing_scenarios}")
            # Dodaj brakujące scenariusze jako błędy
            for missing_id in missing_scenarios:
                missing_scenario = next((s for s in scenarios if s.get('scenario_id') == missing_id), None)
                if missing_scenario:
                    all_detailed_scenarios.append({
                        'scenario_id': missing_id,
                        'test_case_id': f'TC_{len(all_detailed_scenarios) + 1:04d}',
                        'scenario_name': missing_scenario.get('title', 'Brak scenariusza'),
                        'source_sections': missing_scenario.get('source_sections', missing_scenario.get('source_pages', [])),
                        'priority': missing_scenario.get('priority', 'Medium'),
                        'status': 'Error',
                        'steps': [{
                            'step_number': 1,
                            'action': 'Scenariusz nie został przetworzony w etapie 3',
                            'expected_result': 'Wymagana ręczna weryfikacja'
                        }]
                    })
        
        print(f"  Przetworzono {len(all_detailed_scenarios)}/{len(scenarios)} scenariuszy")
        
        # Zapisz szczegółowe scenariusze do pliku Excel
        return self.save_detailed_results(all_detailed_scenarios, results_dir, task_id)
    
    def save_detailed_results(self, detailed_scenarios: List[Dict], results_dir: Path, task_id: str) -> Path:
        """
        Zapisuje szczegółowe scenariusze testowe do pliku Excel z wieloma krokami.
        
        Args:
            detailed_scenarios: Lista szczegółowych scenariuszy z krokami
            results_dir: Katalog wyników
            task_id: ID zadania
            
        Returns:
            Ścieżka do zapisanego pliku
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Scenariusze Testowe"
            
            # Nagłówki
            headers = [
                'Test Case ID',
                'Nazwa scenariusza',
                'Numer kroku',
                'Akcja',
                'Oczekiwany rezultat',
                'Źródło dokumentacji',
                'Priorytet',
                'Status'
            ]
            
            # Styl nagłówków
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Dane - każdy krok w osobnym wierszu
            row_idx = 2
            for scenario in detailed_scenarios:
                test_case_id = scenario.get('test_case_id', '')
                scenario_name = scenario.get('scenario_name', '')
                # Obsługa zarówno source_sections jak i source_pages (dla kompatybilności)
                source_sections = scenario.get('source_sections', [])
                if not source_sections:
                    source_sections = scenario.get('source_pages', [])
                # Upewnij się, że source_sections jest listą
                if not isinstance(source_sections, list):
                    source_sections = []
                # Formatuj sekcje jako czytelny opis (np. "sekcje: Wstęp, Instalacja")
                source_sections_str = self._format_source_sections(source_sections) if source_sections else 'cała dokumentacja'
                priority = scenario.get('priority', 'Medium')
                status = scenario.get('status', 'Draft')
                steps = scenario.get('steps', [])
                if not isinstance(steps, list):
                    steps = []
                steps = self._normalize_steps(steps)
                
                if not steps:
                    # Jeśli brak kroków, dodaj jeden wiersz
                    ws.cell(row=row_idx, column=1, value=test_case_id)
                    ws.cell(row=row_idx, column=2, value=scenario_name)
                    ws.cell(row=row_idx, column=3, value='-')
                    ws.cell(row=row_idx, column=4, value='Brak kroków')
                    ws.cell(row=row_idx, column=5, value='-')
                    ws.cell(row=row_idx, column=6, value=source_sections_str)
                    ws.cell(row=row_idx, column=7, value=priority)
                    ws.cell(row=row_idx, column=8, value=status)
                    row_idx += 1
                else:
                    # Każdy krok w osobnym wierszu
                    for step in steps:
                        ws.cell(row=row_idx, column=1, value=test_case_id)
                        ws.cell(row=row_idx, column=2, value=scenario_name)
                        ws.cell(row=row_idx, column=3, value=step.get('step_number', ''))
                        ws.cell(row=row_idx, column=4, value=step.get('action', ''))
                        ws.cell(row=row_idx, column=5, value=step.get('expected_result', ''))
                        ws.cell(row=row_idx, column=6, value=source_sections_str)
                        ws.cell(row=row_idx, column=7, value=priority)
                        ws.cell(row=row_idx, column=8, value=status)
                        row_idx += 1
            
            # Dostosuj szerokość kolumn
            column_widths = [15, 40, 10, 50, 50, 15, 12, 12]
            for col_idx, width in enumerate(column_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
            
            # Zapisz plik
            results_dir.mkdir(parents=True, exist_ok=True)
            result_file = results_dir / f"wyniki_{task_id}.xlsx"
            wb.save(str(result_file))
            
            print(f"Zapisano {len(detailed_scenarios)} scenariuszy z łączną liczbą {row_idx - 2} kroków do pliku {result_file}")
            return result_file
            
        except ImportError:
            # Fallback: zapisz jako JSON
            results_dir.mkdir(parents=True, exist_ok=True)
            result_file = results_dir / f"wyniki_{task_id}.json"
            
            with open(result_file, 'w', encoding='utf-8') as f:
                json.dump(detailed_scenarios, f, ensure_ascii=False, indent=2)
            
            return result_file