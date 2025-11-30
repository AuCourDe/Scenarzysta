"""
Procesor dokumentów v0.2 - nowy workflow z segmentacją i inteligentnymi opisami obrazów.
Obsługuje formaty: DOCX, PDF, XLSX, TXT
"""
import zipfile
import os
import re
import base64
import requests
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import json
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

# Import ekstraktora dla innych formatów
try:
    from file_extractors import FileExtractor, extract_file
    FILE_EXTRACTORS_AVAILABLE = True
except ImportError:
    FILE_EXTRACTORS_AVAILABLE = False


class ContextLengthError(Exception):
    """Błąd przekroczenia limitu kontekstu/tokenów modelu."""
    pass


class DocumentProcessorV2:
    """
    Procesor dokumentów v0.2 z nowym workflow:
    1. Ekstrakcja tekstu + opisy obrazów przez AI
    2. Segmentacja na fragmenty funkcjonalności
    3. Korelacja plików (opcjonalna)
    4. Generowanie ścieżek testowych
    5. Generowanie szczegółowych scenariuszy
    """
    
    def __init__(self, ollama_url: str = "http://localhost:11434", ollama_model: str = "gemma3:12B", settings: Optional[Dict[str, Any]] = None):
        self.ollama_url = ollama_url
        self.ollama_model = ollama_model
        
        cfg = settings or {}
        self.temperature = float(cfg.get('temperature', 0.2))
        self.top_p = float(cfg.get('top_p', 0.9))
        self.top_k = int(cfg.get('top_k', 40))
        self.max_tokens = int(cfg.get('max_tokens', 8192))
        self.context_length = int(cfg.get('context_length', 16000))
        self.segment_chunk_words = int(cfg.get('segment_chunk_words', 500))
        
        # Śledzenie postępu
        self.processing_stats = {
            'total_chunks': 0,
            'processed_chunks': 0,
            'chunk_times': [],
            'start_time': None,
            'current_stage': 0,
            'total_stages': 4  # Ekstrakcja, Segmentacja, Ścieżki, Scenariusze
        }
        
        # Konfiguracja użytkownika (opcjonalne)
        self.user_config = {
            'custom_paths_description': '',  # Opis wymagań do ścieżek testowych
            'custom_scenarios_description': '',  # Opis wymagań do scenariuszy
            'example_documentation': '',  # Przykład fragmentu dokumentacji
            'example_scenarios': []  # Przykładowe scenariusze użytkownika
        }
    
    def reset_processing_stats(self):
        """Resetuje statystyki przetwarzania."""
        self.processing_stats = {
            'total_chunks': 0,
            'processed_chunks': 0,
            'chunk_times': [],
            'start_time': None,
            'current_stage': 0,
            'total_stages': 4
        }
    
    def reset_user_config(self):
        """Resetuje konfigurację użytkownika do wartości domyślnych."""
        self.user_config = {
            'custom_paths_description': '',
            'custom_scenarios_description': '',
            'example_documentation': '',
            'example_scenarios': []
        }
    
    def set_user_config(self, config: Dict):
        """Ustawia konfigurację użytkownika (opisy, przykłady)."""
        self.user_config.update(config)
    
    def get_dynamic_eta(self) -> Optional[float]:
        """Oblicza dynamiczny ETA na podstawie rzeczywistego postępu."""
        import time
        stats = self.processing_stats
        
        if not stats['chunk_times'] or stats['total_chunks'] == 0:
            return None
        
        avg_chunk_time = sum(stats['chunk_times']) / len(stats['chunk_times'])
        remaining_chunks = stats['total_chunks'] - stats['processed_chunks']
        
        if remaining_chunks <= 0:
            return 0.0
        
        eta = remaining_chunks * avg_chunk_time
        return eta
    
    def _load_prompt(self, prompt_file: str) -> str:
        """Wczytuje prompt z pliku."""
        prompt_path = Path(__file__).parent / prompt_file
        if prompt_path.exists():
            with open(prompt_path, 'r', encoding='utf-8') as f:
                return f.read()
        return ""
    
    def _call_ollama(self, prompt: str, images: List[str] = None, timeout: int = 300) -> str:
        """
        Wywołuje Ollama API.
        
        Args:
            prompt: Prompt tekstowy
            images: Lista base64-encoded obrazów (opcjonalne)
            timeout: Timeout w sekundach
            
        Returns:
            Odpowiedź modelu
        """
        try:
            api_url = f"{self.ollama_url}/api/generate"
            
            payload = {
                "model": self.ollama_model,
                "prompt": prompt,
                "stream": False
            }
            options: Dict[str, Any] = {}
            if self.temperature is not None:
                options["temperature"] = float(self.temperature)
            if self.top_p is not None:
                options["top_p"] = float(self.top_p)
            if self.top_k:
                options["top_k"] = int(self.top_k)
            if self.max_tokens:
                options["num_predict"] = int(self.max_tokens)
            if self.context_length:
                options["num_ctx"] = int(self.context_length)
            if options:
                payload["options"] = options
            
            if images:
                payload["images"] = images
            
            response = requests.post(api_url, json=payload, timeout=timeout)
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', '').strip()
            elif response.status_code == 500:
                error_text = response.text.lower()
                if 'context length' in error_text or 'token' in error_text:
                    raise ContextLengthError(f"Przekroczono limit kontekstu: {response.text[:200]}")
                print(f"Błąd serwera Ollama (500): {response.text[:200]}")
                return ""
            else:
                print(f"Błąd Ollama: {response.status_code} - {response.text[:200]}")
                return ""
                
        except requests.exceptions.Timeout:
            print(f"  ⚠️ Timeout wywołania Ollama ({timeout}s)")
            return ""
        except ContextLengthError:
            raise
        except Exception as e:
            print(f"Błąd wywołania Ollama: {e}")
            return ""
    
    def _call_ollama_with_image(self, prompt: str, image_path: str) -> str:
        """Wywołuje Ollama z obrazem."""
        try:
            if not os.path.exists(image_path):
                return ""
            
            with open(image_path, 'rb') as img_file:
                image_data = base64.b64encode(img_file.read()).decode('utf-8')
            
            api_url = f"{self.ollama_url}/api/chat"
            
            payload = {
                "model": self.ollama_model,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt,
                        "images": [image_data]
                    }
                ],
                "stream": False
            }
            options: Dict[str, Any] = {}
            if self.temperature is not None:
                options["temperature"] = float(self.temperature)
            if self.top_p is not None:
                options["top_p"] = float(self.top_p)
            if self.top_k:
                options["top_k"] = int(self.top_k)
            if self.max_tokens:
                options["num_predict"] = int(self.max_tokens)
            if self.context_length:
                options["num_ctx"] = int(self.context_length)
            if options:
                payload["options"] = options
            
            response = requests.post(api_url, json=payload, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                message = result.get('message', {})
                return message.get('content', '').strip()
            return ""
            
        except Exception as e:
            print(f"  ⚠️ Błąd analizy obrazu: {e}")
            return ""

    # =========================================================================
    # ETAP 0: EKSTRAKCJA TEKSTU Z OPISAMI OBRAZÓW
    # =========================================================================
    
    def extract_and_describe(self, file_path: str, output_dir: str) -> Dict:
        """
        NOWY WORKFLOW: Ekstrakcja tekstu z dokumentu + opisy wszystkich obrazów/tabel przez AI.
        
        Zwraca spójny tekst z wstawionymi opisami obrazów w miejscach gdzie były grafiki.
        """
        print(f"[EKSTRAKCJA] Rozpoczynam ekstrakcję z: {Path(file_path).name}")
        
        ext = Path(file_path).suffix.lower()
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # Ekstrakcja zależna od formatu
        if ext == '.docx':
            extracted = self._extract_docx_with_positions(file_path, output_dir)
        elif ext == '.pdf':
            extracted = self._extract_pdf_with_positions(file_path, output_dir)
        elif ext in ['.xlsx', '.xls']:
            extracted = self._extract_excel(file_path, output_dir)
        elif ext == '.txt':
            extracted = self._extract_txt(file_path)
        else:
            raise ValueError(f"Nieobsługiwany format: {ext}")
        
        # Opisz wszystkie obrazy przez AI
        print(f"[EKSTRAKCJA] Opisuję {len(extracted.get('images', []))} obrazów przez AI...")
        image_descriptions = self._describe_all_images(extracted.get('images', []))
        
        # Połącz tekst z opisami obrazów
        combined_text = self._combine_text_with_image_descriptions(
            extracted.get('text_with_placeholders', ''),
            image_descriptions
        )
        
        # Zapisz połączony tekst
        combined_file = output_path / "dokument_z_opisami.txt"
        with open(combined_file, 'w', encoding='utf-8') as f:
            f.write(combined_text)
        
        print(f"[EKSTRAKCJA] Zapisano połączony dokument: {combined_file.name}")
        print(f"[EKSTRAKCJA] Długość dokumentu: {len(combined_text)} znaków, {len(combined_text.split())} słów")
        
        return {
            'combined_text': combined_text,
            'combined_text_file': str(combined_file),
            'images': extracted.get('images', []),
            'image_descriptions': image_descriptions,
            'metadata': extracted.get('metadata', {})
        }
    
    def _extract_docx_with_positions(self, docx_path: str, output_dir: str) -> Dict:
        """Ekstrakcja z DOCX z zachowaniem pozycji obrazów."""
        output_path = Path(output_dir)
        images_dir = output_path / "images"
        images_dir.mkdir(exist_ok=True)
        
        doc = Document(docx_path)
        
        # Ekstrakcja obrazów
        images = []
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            for file_info in zip_ref.namelist():
                if file_info.startswith('word/media/'):
                    image_data = zip_ref.read(file_info)
                    image_filename = os.path.basename(file_info)
                    image_path = images_dir / image_filename
                    
                    with open(image_path, 'wb') as img_file:
                        img_file.write(image_data)
                    
                    images.append({
                        'filename': image_filename,
                        'path': str(image_path),
                        'original_path': file_info
                    })
        
        # Mapa relacji obrazów
        image_rels = {}
        for rel in doc.part.rels.values():
            if "image" in getattr(rel, 'reltype', ''):
                try:
                    image_filename = os.path.basename(rel.target_ref)
                    image_rels[rel.rId] = image_filename
                except:
                    pass
        
        # Ekstrakcja tekstu z placeholderami obrazów
        text_parts = []
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            
            # Sprawdź czy paragraf zawiera obrazy
            for run in para.runs:
                blips = run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and rId in image_rels:
                        image_filename = image_rels[rId]
                        placeholder = f"\n[__IMAGE__{image_filename}__]\n"
                        para_text += placeholder
            
            if para_text:
                text_parts.append(para_text)
        
        # Ekstrakcja tabel
        for table_idx, table in enumerate(doc.tables):
            table_text = f"\n[__TABLE__{table_idx}__]\n"
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(row_cells) + "\n"
            text_parts.append(table_text)
        
        text_with_placeholders = "\n\n".join(text_parts)
        
        return {
            'text_with_placeholders': text_with_placeholders,
            'images': images,
            'metadata': {
                'filename': os.path.basename(docx_path),
                'total_images': len(images),
                'format': 'docx'
            }
        }
    
    def _extract_pdf_with_positions(self, pdf_path: str, output_dir: str) -> Dict:
        """Ekstrakcja z PDF z zachowaniem pozycji obrazów."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            # Fallback bez obrazów
            return self._extract_pdf_text_only(pdf_path)
        
        output_path = Path(output_dir)
        images_dir = output_path / "images"
        images_dir.mkdir(exist_ok=True)
        
        doc = fitz.open(pdf_path)
        
        text_parts = []
        images = []
        
        for page_num, page in enumerate(doc):
            page_text = f"[Strona {page_num + 1}]\n"
            page_text += page.get_text()
            
            # Ekstrakcja obrazów ze strony
            image_list = page.get_images()
            for img_idx, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image_filename = f"page{page_num + 1}_img{img_idx}.{image_ext}"
                    image_path = images_dir / image_filename
                    
                    with open(image_path, 'wb') as f:
                        f.write(image_bytes)
                    
                    images.append({
                        'filename': image_filename,
                        'path': str(image_path),
                        'page': page_num + 1
                    })
                    
                    # Dodaj placeholder
                    page_text += f"\n[__IMAGE__{image_filename}__]\n"
                except Exception as e:
                    print(f"  Błąd ekstrakcji obrazu: {e}")
            
            text_parts.append(page_text)
        
        doc.close()
        
        return {
            'text_with_placeholders': "\n\n".join(text_parts),
            'images': images,
            'metadata': {
                'filename': os.path.basename(pdf_path),
                'total_images': len(images),
                'page_count': len(text_parts),
                'format': 'pdf'
            }
        }
    
    def _extract_pdf_text_only(self, pdf_path: str) -> Dict:
        """Fallback dla PDF bez PyMuPDF."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text_parts.append(f"[Strona {page_num + 1}]\n{text}")
            
            return {
                'text_with_placeholders': "\n\n".join(text_parts),
                'images': [],
                'metadata': {'filename': os.path.basename(pdf_path), 'format': 'pdf'}
            }
        except ImportError:
            raise ImportError("Zainstaluj PyMuPDF lub pdfplumber")
    
    def _extract_excel(self, file_path: str, output_dir: str) -> Dict:
        """Ekstrakcja z Excel."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"\n[Arkusz: {sheet_name}]\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return {
                'text_with_placeholders': "\n".join(text_parts),
                'images': [],
                'metadata': {'filename': os.path.basename(file_path), 'format': 'xlsx'}
            }
        except ImportError:
            raise ImportError("Zainstaluj openpyxl")
    
    def _extract_txt(self, file_path: str) -> Dict:
        """Ekstrakcja z TXT."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return {
            'text_with_placeholders': text,
            'images': [],
            'metadata': {'filename': os.path.basename(file_path), 'format': 'txt'}
        }
    
    def _describe_all_images(self, images: List[Dict]) -> Dict[str, str]:
        """Opisuje wszystkie obrazy przez AI."""
        descriptions = {}
        prompt_template = self._load_prompt('prompt_images.txt')
        
        # Filtruj małe obrazy (prawdopodobnie ikony/loga)
        filtered_images = []
        for img in images:
            try:
                with Image.open(img['path']) as pil_img:
                    w, h = pil_img.size
                    if w >= 50 and h >= 50:
                        filtered_images.append(img)
                    else:
                        print(f"  Pomijam mały obraz: {img['filename']} ({w}x{h})")
            except:
                filtered_images.append(img)
        
        for idx, img in enumerate(filtered_images, 1):
            print(f"  Opisuję obraz {idx}/{len(filtered_images)}: {img['filename']}")
            
            description = self._call_ollama_with_image(prompt_template, img['path'])
            
            if description:
                descriptions[img['filename']] = description
            else:
                descriptions[img['filename']] = f"[W tym miejscu była grafika: {img['filename']}]"
        
        return descriptions
    
    def _combine_text_with_image_descriptions(self, text: str, descriptions: Dict[str, str]) -> str:
        """Łączy tekst z opisami obrazów."""
        combined = text
        
        for filename, description in descriptions.items():
            placeholder = f"[__IMAGE__{filename}__]"
            combined = combined.replace(placeholder, f"\n{description}\n")
        
        # Usuń pozostałe nieopisane placeholdery
        combined = re.sub(r'\[__IMAGE__[^\]]+__\]', '', combined)
        combined = re.sub(r'\[__TABLE__\d+__\]', '', combined)
        
        return combined.strip()

    # =========================================================================
    # ETAP 1: SEGMENTACJA DOKUMENTU
    # =========================================================================
    
    def segment_document(self, combined_text: str, processing_dir: Path, correlate: bool = False) -> List[Dict]:
        """
        SEGMENTACJA: Dzieli dokument na fragmenty po ~500 słów i analizuje każdy fragment.
        
        Returns:
            Lista segmentów z metadanymi o funkcjonalnościach
        """
        import time
        
        print(f"[SEGMENTACJA] Rozpoczynam segmentację dokumentu...")
        
        # Podziel na fragmenty po zadanej liczbie słów (domyślnie ~500)
        words = combined_text.split()
        chunk_size = max(100, self.segment_chunk_words or 500)
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append({
                'index': len(chunks) + 1,
                'text': chunk_text,
                'word_count': len(chunk_words)
            })
        
        print(f"[SEGMENTACJA] Podzielono na {len(chunks)} fragmentów po ~{chunk_size} słów")
        
        # Wczytaj prompt segmentacji
        prompt_template = self._load_prompt('prompt_segmentation.txt')
        
        # Dodaj fragment o korelacji jeśli włączona
        if correlate:
            prompt_template += "\n\nDODATKOWO: Szczególnie zwróć uwagę na pola needs_correlation - zaznacz true jeśli fragment odwołuje się do informacji, które mogą być w innym dokumencie."
        
        # Analizuj każdy fragment
        self.processing_stats['current_stage'] = 1
        self.processing_stats['total_chunks'] = len(chunks)
        self.processing_stats['processed_chunks'] = 0
        self.processing_stats['start_time'] = time.time()
        
        analyzed_chunks = []
        
        for chunk in chunks:
            chunk_start = time.time()
            
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Analizuję fragment {chunk['index']}/{len(chunks)}...{eta_str}")
            
            full_prompt = f"{prompt_template}\n\nFRAGMENT {chunk['index']} ({chunk['word_count']} słów):\n{chunk['text']}"
            
            response = self._call_ollama(full_prompt)
            
            # Parsuj JSON
            analysis = self._parse_segmentation_response(response, chunk['index'])
            analysis['original_text'] = chunk['text']
            analysis['word_count'] = chunk['word_count']
            
            analyzed_chunks.append(analysis)
            
            # Aktualizuj statystyki
            chunk_time = time.time() - chunk_start
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = chunk['index']
        
        # Utwórz podsumowanie wszystkich fragmentów
        summary = self._create_document_summary(analyzed_chunks)
        
        # Podziel oryginalny tekst na logiczne fragmenty
        logical_segments = self._create_logical_segments(combined_text, analyzed_chunks, summary)
        
        # Zapisz wyniki segmentacji
        self._save_segmentation_results(processing_dir, analyzed_chunks, summary, logical_segments)
        
        print(f"[SEGMENTACJA] Utworzono {len(logical_segments)} logicznych segmentów")
        
        return logical_segments
    
    def _parse_segmentation_response(self, response: str, fragment_num: int) -> Dict:
        """Parsuje odpowiedź segmentacji."""
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
        except:
            pass
        
        # Fallback
        return {
            'fragment_number': fragment_num,
            'topic': f'Fragment {fragment_num}',
            'topic_type': 'unknown',
            'summary': 'Nie udało się przeanalizować fragmentu',
            'start_sentence': None,
            'end_sentence': None,
            'prerequisites': [],
            'related_topics': [],
            'needs_correlation': False
        }
    
    def _create_document_summary(self, analyzed_chunks: List[Dict]) -> Dict:
        """Tworzy podsumowanie całego dokumentu na podstawie analizy fragmentów."""
        topics = []
        all_prerequisites = set()
        correlation_needed = []
        
        for chunk in analyzed_chunks:
            topic = chunk.get('topic', '')
            if topic and topic not in [t['name'] for t in topics]:
                topics.append({
                    'name': topic,
                    'type': chunk.get('topic_type', 'unknown'),
                    'start_sentence': chunk.get('start_sentence'),
                    'end_sentence': chunk.get('end_sentence'),
                    'fragments': [chunk.get('fragment_number')]
                })
            elif topic:
                # Dodaj fragment do istniejącego tematu
                for t in topics:
                    if t['name'] == topic:
                        t['fragments'].append(chunk.get('fragment_number'))
                        if chunk.get('end_sentence'):
                            t['end_sentence'] = chunk.get('end_sentence')
            
            for prereq in chunk.get('prerequisites', []):
                all_prerequisites.add(prereq)
            
            if chunk.get('needs_correlation'):
                correlation_needed.append(chunk.get('fragment_number'))
        
        return {
            'topics': topics,
            'total_topics': len(topics),
            'all_prerequisites': list(all_prerequisites),
            'fragments_needing_correlation': correlation_needed
        }
    
    def _create_logical_segments(self, original_text: str, analyzed_chunks: List[Dict], summary: Dict) -> List[Dict]:
        """
        Tworzy logiczne segmenty dokumentu na podstawie analizy.
        Każdy segment zawiera pełną treść tematu + wymagania wstępne.
        """
        segments = []
        
        for topic in summary.get('topics', []):
            topic_name = topic.get('name', 'Nieznany temat')
            fragment_indices = topic.get('fragments', [])
            
            # Zbierz tekst ze wszystkich fragmentów tego tematu
            topic_text = []
            topic_prerequisites = set()
            
            for chunk in analyzed_chunks:
                if chunk.get('fragment_number') in fragment_indices:
                    original_chunk_text = chunk.get('original_text', '')
                    if original_chunk_text:
                        span = self._extract_segment_span(
                            original_chunk_text,
                            chunk.get('start_sentence'),
                            chunk.get('end_sentence')
                        )
                        topic_text.append(span)
                    for prereq in chunk.get('prerequisites', []):
                        topic_prerequisites.add(prereq)
            
            # Dodaj wymagania wstępne na początku segmentu
            prerequisites_text = ""
            if topic_prerequisites:
                prerequisites_text = "WYMAGANIA WSTĘPNE:\n"
                for prereq in topic_prerequisites:
                    prerequisites_text += f"- {prereq}\n"
                prerequisites_text += "\n"
            
            segment = {
                'segment_id': f"SEG_{len(segments) + 1:03d}",
                'topic': topic_name,
                'prerequisites': list(topic_prerequisites),
                'prerequisites_text': prerequisites_text,
                'content': '\n\n'.join(topic_text),
                'full_text': prerequisites_text + '\n\n'.join(topic_text),
                'fragments': fragment_indices,
                'needs_correlation': any(
                    chunk.get('needs_correlation') 
                    for chunk in analyzed_chunks 
                    if chunk.get('fragment_number') in fragment_indices
                )
            }
            
            segments.append(segment)
        
        # Jeśli nie udało się wyodrębnić tematów, użyj całego tekstu jako jednego segmentu
        if not segments:
            segments.append({
                'segment_id': 'SEG_001',
                'topic': 'Cały dokument',
                'prerequisites': [],
                'prerequisites_text': '',
                'content': original_text,
                'full_text': original_text,
                'fragments': list(range(1, len(analyzed_chunks) + 1)),
                'needs_correlation': False
            })
        
        return segments
    
    def _extract_segment_span(self, original_text: str, start_sentence, end_sentence) -> str:
        if not original_text:
            return ""
        text = original_text
        start_idx = 0
        end_idx = len(text)
        
        if start_sentence:
            start_str = str(start_sentence).strip()
            if start_str:
                idx = text.find(start_str)
                if idx != -1:
                    start_idx = idx
        
        if end_sentence:
            end_str = str(end_sentence).strip()
            if end_str:
                idx = text.rfind(end_str)
                if idx != -1 and idx >= start_idx:
                    end_idx = idx + len(end_str)
        
        span = text[start_idx:end_idx].strip()
        if not span:
            return original_text
        return span
    
    def _save_segmentation_results(self, processing_dir: Path, chunks: List[Dict], summary: Dict, segments: List[Dict]):
        """Zapisuje wyniki segmentacji."""
        processing_dir.mkdir(parents=True, exist_ok=True)
        
        # Zapisz analizę fragmentów
        with open(processing_dir / "analiza_fragmentow.json", 'w', encoding='utf-8') as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)
        
        # Zapisz podsumowanie
        with open(processing_dir / "podsumowanie_dokumentu.json", 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        
        # Zapisz każdy segment jako osobny plik
        segments_dir = processing_dir / "segmenty"
        segments_dir.mkdir(exist_ok=True)
        
        for segment in segments:
            segment_file = segments_dir / f"{segment['segment_id']}_{self._sanitize_filename(segment['topic'])}.txt"
            with open(segment_file, 'w', encoding='utf-8') as f:
                f.write(f"# {segment['topic']}\n\n")
                f.write(segment['full_text'])
    
    def _sanitize_filename(self, name: str) -> str:
        """Czyści nazwę pliku."""
        return re.sub(r'[^\w\s-]', '', name)[:50].strip().replace(' ', '_')

    # =========================================================================
    # ETAP 2: KORELACJA DOKUMENTÓW (opcjonalna)
    # =========================================================================
    
    def correlate_segments(self, all_segments: Dict[str, List[Dict]], processing_dir: Path) -> List[Dict]:
        """
        Koreluje segmenty z różnych dokumentów.
        
        Args:
            all_segments: Dict {filename: [segments]}
            
        Returns:
            Lista skorelowanych grup segmentów
        """
        print(f"[KORELACJA] Rozpoczynam korelację {len(all_segments)} dokumentów...")
        
        # Zbierz wszystkie podsumowania
        all_summaries = []
        for filename, segments in all_segments.items():
            for seg in segments:
                all_summaries.append({
                    'filename': filename,
                    'segment_id': seg['segment_id'],
                    'topic': seg['topic'],
                    'prerequisites': seg.get('prerequisites', []),
                    'needs_correlation': seg.get('needs_correlation', False)
                })
        
        # Przygotuj prompt do korelacji
        summaries_json = json.dumps(all_summaries, ensure_ascii=False, indent=2)
        
        prompt = f"""Przeanalizuj poniższe segmenty z różnych dokumentów i znajdź powiązania między nimi.

SEGMENTY:
{summaries_json}

Zwróć JSON w formacie:
{{
    "correlated_groups": [
        {{
            "group_id": "CORR_001",
            "topic": "wspólny temat",
            "segments": [
                {{"filename": "plik1.docx", "segment_id": "SEG_001"}},
                {{"filename": "plik2.docx", "segment_id": "SEG_003"}}
            ],
            "reason": "dlaczego te segmenty są powiązane"
        }}
    ]
}}

Zwracaj TYLKO JSON, bez tekstu przed/po."""

        response = self._call_ollama(prompt)
        
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group(0))
                
                # Zapisz wyniki korelacji
                with open(processing_dir / "korelacja_dokumentow.json", 'w', encoding='utf-8') as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
                
                return result.get('correlated_groups', [])
        except:
            pass
        
        return []

    # =========================================================================
    # ETAP 3: GENEROWANIE ŚCIEŻEK TESTOWYCH (nowy prompt1)
    # =========================================================================
    
    def generate_test_paths(self, segments: List[Dict], processing_dir: Path, results_dir: Path, task_id: str) -> List[Dict]:
        """
        Generuje ścieżki testowe dla każdego segmentu dokumentacji.
        """
        import time
        
        print(f"[ŚCIEŻKI] Generuję ścieżki testowe dla {len(segments)} segmentów...")
        
        prompt_template = self._load_prompt('prompt_paths.txt')
        
        # Dodaj customowy opis użytkownika jeśli jest
        user_paths_desc = self.user_config.get('custom_paths_description', '')
        if user_paths_desc:
            print(f"[ŚCIEŻKI] Dodaję wymagania użytkownika do promptu: {user_paths_desc[:80]}...")
            prompt_template += f"\n\nDODATKOWE WYMAGANIA OD UŻYTKOWNIKA:\n{user_paths_desc}"
        
        # Dodaj przykład użytkownika jeśli jest
        example_doc = self.user_config.get('example_documentation', '')
        example_scenarios = self.user_config.get('example_scenarios', [])
        if example_doc and example_scenarios:
            prompt_template += f"\n\nPRZYKŁAD OD UŻYTKOWNIKA:\nDokumentacja:\n{example_doc}\n\nPrzykładowe scenariusze:\n{json.dumps(example_scenarios, ensure_ascii=False)}"
        
        self.processing_stats['current_stage'] = 2
        self.processing_stats['total_chunks'] = len(segments)
        self.processing_stats['processed_chunks'] = 0
        
        all_paths = []
        
        for idx, segment in enumerate(segments, 1):
            chunk_start = time.time()
            
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Generuję ścieżki dla segmentu {idx}/{len(segments)}: {segment['topic']}...{eta_str}")
            
            full_prompt = f"""{prompt_template}

SEGMENT DOKUMENTACJI ({segment['segment_id']} - {segment['topic']}):

{segment.get('prerequisites_text', '')}

{segment['content']}

PAMIĘTAJ: 
- Wygeneruj WSZYSTKIE możliwe ścieżki pozytywne (happy path)
- Wygeneruj WSZYSTKIE możliwe ścieżki negatywne
- Wygeneruj przypadki brzegowe
- Zwracasz TYLKO JSON"""

            response = self._call_ollama(full_prompt)
            
            # Parsuj odpowiedź
            paths = self._parse_paths_response(response, segment)
            
            # Zapisz ścieżki dla tego segmentu
            segment_paths_file = processing_dir / f"sciezki_{segment['segment_id']}.json"
            with open(segment_paths_file, 'w', encoding='utf-8') as f:
                json.dump(paths, f, ensure_ascii=False, indent=2)
            
            # Zapisz każdą ścieżkę osobno
            paths_dir = processing_dir / "sciezki"
            paths_dir.mkdir(exist_ok=True)
            
            for path_idx, path in enumerate(paths, 1):
                path_file = paths_dir / f"{segment['segment_id']}_sciezka_{path_idx}.txt"
                with open(path_file, 'w', encoding='utf-8') as f:
                    f.write(f"ID: {path.get('id', '')}\n")
                    f.write(f"Typ: {path.get('type', '')}\n")
                    f.write(f"Tytuł: {path.get('title', '')}\n")
                    f.write(f"Opis: {path.get('description', '')}\n")
                    f.write(f"Segment źródłowy: {segment['segment_id']} - {segment['topic']}\n")
            
            all_paths.extend(paths)
            
            chunk_time = time.time() - chunk_start
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = idx
        
        # Zapisz wszystkie ścieżki
        results_dir.mkdir(parents=True, exist_ok=True)
        all_paths_file = results_dir / f"etap1_sciezki_testowe_{task_id}.json"
        with open(all_paths_file, 'w', encoding='utf-8') as f:
            json.dump(all_paths, f, ensure_ascii=False, indent=2)
        
        print(f"[ŚCIEŻKI] Wygenerowano łącznie {len(all_paths)} ścieżek testowych")
        
        return all_paths
    
    def _parse_paths_response(self, response: str, segment: Dict) -> List[Dict]:
        """Parsuje odpowiedź z ścieżkami testowymi."""
        paths = []
        
        try:
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                raw_paths = json.loads(json_match.group(0))
                
                for idx, path in enumerate(raw_paths, 1):
                    if isinstance(path, dict):
                        path['id'] = path.get('id', f"PATH_{len(paths) + 1:03d}")
                        path['source_segment'] = segment['segment_id']
                        path['source_topic'] = segment['topic']
                        path['prerequisites'] = segment.get('prerequisites', [])
                        paths.append(path)
                    elif isinstance(path, str):
                        paths.append({
                            'id': f"PATH_{len(paths) + 1:03d}",
                            'title': path,
                            'description': path,
                            'type': 'happy_path',
                            'source_segment': segment['segment_id'],
                            'source_topic': segment['topic'],
                            'prerequisites': segment.get('prerequisites', [])
                        })
        except json.JSONDecodeError:
            print(f"  ⚠️ Błąd parsowania JSON dla segmentu {segment['segment_id']}")
        
        if not paths:
            paths.append({
                'id': 'PATH_001',
                'title': f"Ścieżka dla: {segment['topic']}",
                'description': 'Wymaga ręcznej weryfikacji',
                'type': 'happy_path',
                'source_segment': segment['segment_id'],
                'source_topic': segment['topic'],
                'prerequisites': segment.get('prerequisites', [])
            })
        
        return paths

    # =========================================================================
    # ETAP 4: GENEROWANIE SZCZEGÓŁOWYCH SCENARIUSZY (nowy prompt2)
    # =========================================================================
    
    def generate_detailed_scenarios(self, paths: List[Dict], segments: List[Dict], processing_dir: Path, results_dir: Path, task_id: str, progress_callback=None) -> Path:
        """
        Generuje szczegółowe scenariusze testowe ze ścieżek.
        Każda ścieżka + jej segment dokumentacji -> szczegółowy scenariusz.
        
        Args:
            progress_callback: Opcjonalna funkcja callback(current, total) do raportowania postępu
        """
        import time
        
        print(f"[SCENARIUSZE] Generuję szczegółowe scenariusze dla {len(paths)} ścieżek...")
        
        prompt_template = self._load_prompt('prompt_scenario.txt')
        if not prompt_template:
            prompt_template = self._get_default_scenario_prompt()
        
        # Dodaj customowy opis użytkownika
        user_scenario_desc = self.user_config.get('custom_scenarios_description', '')
        if user_scenario_desc:
            print(f"[SCENARIUSZE] Dodaję wymagania użytkownika do promptu: {user_scenario_desc[:80]}...")
            prompt_template += f"\n\nDODATKOWE WYMAGANIA OD UŻYTKOWNIKA:\n{user_scenario_desc}"
        
        # Mapa segment_id -> segment
        segment_map = {seg['segment_id']: seg for seg in segments}
        
        self.processing_stats['current_stage'] = 3
        self.processing_stats['total_chunks'] = len(paths)
        self.processing_stats['processed_chunks'] = 0
        
        all_scenarios = []
        global_scenario_counter = 0  # Globalny licznik scenariuszy
        
        for idx, path in enumerate(paths, 1):
            chunk_start = time.time()
            
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Generuję scenariusz {idx}/{len(paths)}: {path.get('title', '')[:50]}...{eta_str}")
            
            # Pobierz segment źródłowy
            segment_id = path.get('source_segment', '')
            segment = segment_map.get(segment_id, {})
            segment_content = segment.get('full_text', segment.get('content', ''))
            
            # Przygotuj prompt
            path_json = json.dumps(path, ensure_ascii=False, indent=2)
            
            full_prompt = f"""{prompt_template}

ŚCIEŻKA TESTOWA:
{path_json}

FRAGMENT DOKUMENTACJI:
{segment_content[:8000]}

WYMAGANIA WSTĘPNE (z dokumentacji):
{', '.join(path.get('prerequisites', [])) or 'Brak'}

PAMIĘTAJ:
- Scenariusz MUSI zawierać DUŻO szczegółowych kroków (minimum 5-10)
- Kroki bazują TYLKO na dokumentacji, nie na wiedzy ogólnej
- Każdy krok: akcja + oczekiwany rezultat
- Wszystko po POLSKU
- Zwracasz TYLKO JSON"""

            response = self._call_ollama(full_prompt, timeout=180)
            
            # Parsuj scenariusz - używamy globalnego licznika dla unikalnych ID
            global_scenario_counter += 1
            scenario = self._parse_scenario_response(response, path, global_scenario_counter)
            all_scenarios.append(scenario)
            
            chunk_time = time.time() - chunk_start
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = idx
            
            # Raportuj postęp przez callback
            if progress_callback:
                progress_callback(idx, len(paths))
        
        # Zapisz do Excel
        results_dir.mkdir(parents=True, exist_ok=True)
        result_file = self._save_scenarios_to_excel(all_scenarios, results_dir, task_id)
        
        # Zapisz też JSON
        json_file = results_dir / f"etap2_scenariusze_{task_id}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(all_scenarios, f, ensure_ascii=False, indent=2)
        
        print(f"[SCENARIUSZE] Wygenerowano {len(all_scenarios)} scenariuszy")
        
        return result_file
    
    def _get_default_scenario_prompt(self) -> str:
        """Domyślny prompt do generowania scenariuszy."""
        return """System: Jesteś ekspertem QA tworzącym szczegółowe scenariusze testowe.

ZADANIE: Na podstawie ścieżki testowej i fragmentu dokumentacji utwórz szczegółowy scenariusz testowy.

WYMAGANIA:
- Scenariusz musi zawierać WIELE kroków (minimum 5-10)
- Każdy krok musi być szczegółowy i wykonalny
- Kroki bazują TYLKO na dokumentacji, nie na wiedzy ogólnej
- Wszystko w języku POLSKIM

FORMAT WYJŚCIOWY (TYLKO JSON):
{
    "scenario_id": "SCEN_XXX",
    "test_case_id": "TC_XXXX",
    "path_type": "positive/negative/edge_case",
    "test_path": "nazwa ścieżki testowej",
    "scenario_title": "tytuł scenariusza",
    "prerequisites": ["wymaganie 1", "wymaganie 2"],
    "documentation_section": "nazwa sekcji dokumentacji",
    "steps": [
        {
            "step_number": 1,
            "action": "szczegółowa akcja do wykonania",
            "expected_result": "oczekiwany rezultat"
        }
    ]
}"""
    
    def _parse_scenario_response(self, response: str, path: Dict, scenario_num: int) -> Dict:
        """Parsuje odpowiedź ze scenariuszem. scenario_num to globalny, unikalny numer scenariusza."""
        try:
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                scenario = json.loads(json_match.group(0))
                
                # ZAWSZE nadpisuj ID - używaj globalnego licznika dla unikalności
                scenario['scenario_id'] = f"SCEN_{scenario_num:03d}"
                scenario['test_case_id'] = f"TC_{scenario_num:04d}"
                
                # Uzupełnij pozostałe pola
                scenario['path_type'] = scenario.get('path_type', path.get('type', 'positive'))
                scenario['test_path'] = scenario.get('test_path', path.get('title', ''))
                scenario['scenario_title'] = scenario.get('scenario_title', path.get('title', ''))
                scenario['prerequisites'] = scenario.get('prerequisites', path.get('prerequisites', []))
                scenario['documentation_section'] = scenario.get('documentation_section', path.get('source_topic', ''))
                
                # Upewnij się że są kroki
                if 'steps' not in scenario or not scenario['steps']:
                    scenario['steps'] = [{
                        'step_number': 1,
                        'action': 'Wymaga ręcznego uzupełnienia',
                        'expected_result': 'Zgodnie z dokumentacją'
                    }]
                
                return scenario
        except:
            pass
        
        # Fallback
        return {
            'scenario_id': f"SCEN_{scenario_num:03d}",
            'test_case_id': f"TC_{scenario_num:04d}",
            'path_type': path.get('type', 'positive'),
            'test_path': path.get('title', ''),
            'scenario_title': path.get('title', ''),
            'prerequisites': path.get('prerequisites', []),
            'documentation_section': path.get('source_topic', ''),
            'steps': [{
                'step_number': 1,
                'action': 'Wymaga ręcznego uzupełnienia na podstawie dokumentacji',
                'expected_result': 'Zgodnie z dokumentacją'
            }]
        }
    
    def _save_scenarios_to_excel(self, scenarios: List[Dict], results_dir: Path, task_id: str) -> Path:
        """Zapisuje scenariusze do pliku Excel."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Scenariusze Testowe"
            
            # Nagłówki zgodne z nowym workflow
            headers = [
                'Test Case ID',
                'Typ ścieżki',
                'Ścieżka testowa',
                'Tytuł scenariusza',
                'Numer kroku',
                'Akcja',
                'Oczekiwany rezultat',
                'Wymagania wstępne',
                'Sekcja dokumentacji'
            ]
            
            # Styl nagłówków
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = thin_border
            
            # Dane
            row_idx = 2
            for scenario in scenarios:
                test_case_id = scenario.get('test_case_id', '')
                path_type = scenario.get('path_type', '')
                test_path = scenario.get('test_path', '')
                scenario_title = scenario.get('scenario_title', '')
                prerequisites = ', '.join(scenario.get('prerequisites', []))
                doc_section = scenario.get('documentation_section', '')
                steps = scenario.get('steps', [])
                
                if not steps:
                    steps = [{'step_number': 1, 'action': 'Brak kroków', 'expected_result': '-'}]
                
                for step in steps:
                    ws.cell(row=row_idx, column=1, value=test_case_id).border = thin_border
                    ws.cell(row=row_idx, column=2, value=path_type).border = thin_border
                    ws.cell(row=row_idx, column=3, value=test_path).border = thin_border
                    ws.cell(row=row_idx, column=4, value=scenario_title).border = thin_border
                    ws.cell(row=row_idx, column=5, value=step.get('step_number', '')).border = thin_border
                    
                    action_cell = ws.cell(row=row_idx, column=6, value=step.get('action', ''))
                    action_cell.alignment = Alignment(wrap_text=True)
                    action_cell.border = thin_border
                    
                    result_cell = ws.cell(row=row_idx, column=7, value=step.get('expected_result', ''))
                    result_cell.alignment = Alignment(wrap_text=True)
                    result_cell.border = thin_border
                    
                    ws.cell(row=row_idx, column=8, value=prerequisites).border = thin_border
                    ws.cell(row=row_idx, column=9, value=doc_section).border = thin_border
                    
                    row_idx += 1
            
            # Szerokości kolumn
            column_widths = [15, 12, 35, 35, 8, 50, 50, 30, 25]
            for col_idx, width in enumerate(column_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
            
            # Zapisz
            result_file = results_dir / f"wyniki_{task_id}.xlsx"
            wb.save(str(result_file))
            
            print(f"  Zapisano {len(scenarios)} scenariuszy ({row_idx - 2} kroków) do: {result_file.name}")
            return result_file
            
        except ImportError:
            # Fallback do JSON
            result_file = results_dir / f"wyniki_{task_id}.json"
            with open(result_file, 'w', encoding='utf-8') as f:
                json.dump(scenarios, f, ensure_ascii=False, indent=2)
            return result_file
    
    # =========================================================================
    # ETAP 5: GENEROWANIE SZABLONÓW TESTÓW AUTOMATYCZNYCH (v0.4)
    # =========================================================================
    
    def generate_automation_tests(self, scenarios_excel_path: Path, results_dir: Path, 
                                   task_id: str, automation_config: Dict = None,
                                   progress_callback=None) -> Path:
        """
        Generuje szablony testów automatycznych na podstawie scenariuszy manualnych.
        
        Args:
            scenarios_excel_path: Ścieżka do pliku Excel ze scenariuszami manualnymi
            results_dir: Katalog wynikowy
            task_id: ID zadania
            automation_config: Konfiguracja (custom_prompt, example_files)
            progress_callback: Callback do raportowania postępu
            
        Returns:
            Ścieżka do ZIP z wygenerowanymi plikami Java
        """
        import time
        import zipfile
        
        print(f"[AUTOMATYZACJA] Generuję szablony testów automatycznych...")
        
        # Wczytaj scenariusze z Excel
        scenarios = self._load_scenarios_from_excel(scenarios_excel_path)
        print(f"[AUTOMATYZACJA] Wczytano {len(scenarios)} scenariuszy")
        
        # Załaduj prompt
        config = automation_config or {}
        custom_prompt = config.get('custom_prompt', '')
        example_files = config.get('example_files', [])
        
        if custom_prompt or example_files:
            # Użyj custom promptu użytkownika
            prompt_template = custom_prompt or "Wygeneruj klasę testową w Javie na podstawie scenariusza testowego."
            if example_files:
                prompt_template += "\n\nPRZYKŁADY KODU OD UŻYTKOWNIKA:\n"
                for ef in example_files:
                    prompt_template += f"\n--- {ef['filename']} ---\n{ef['content']}\n"
        else:
            # Użyj domyślnego promptu
            prompt_template = self._load_prompt('prompt_automation.txt')
            if not prompt_template:
                prompt_template = self._get_default_automation_prompt()
        
        self.processing_stats['current_stage'] = 4
        self.processing_stats['total_chunks'] = len(scenarios)
        self.processing_stats['processed_chunks'] = 0
        
        # Katalog na pliki Java
        automation_dir = results_dir / f"automation_{task_id}"
        automation_dir.mkdir(parents=True, exist_ok=True)
        
        generated_files = []
        
        for idx, scenario in enumerate(scenarios, 1):
            chunk_start = time.time()
            
            eta = self.get_dynamic_eta()
            eta_str = f" | ETA: {int(eta)}s" if eta else ""
            print(f"  Generuję test {idx}/{len(scenarios)}: {scenario.get('title', '')[:40]}...{eta_str}")
            
            # Przygotuj prompt
            scenario_text = self._format_scenario_for_automation(scenario)
            full_prompt = f"{prompt_template}\n\n{scenario_text}"
            
            response = self._call_ollama(full_prompt, timeout=180)
            
            # Wyodrębnij kod Java
            java_code = self._extract_java_code(response, scenario)
            
            # Zapisz plik Java
            class_name = self._generate_class_name(scenario.get('title', f'Test{idx}'))
            java_file = automation_dir / f"{class_name}.java"
            with open(java_file, 'w', encoding='utf-8') as f:
                f.write(java_code)
            
            generated_files.append(java_file)
            
            chunk_time = time.time() - chunk_start
            self.processing_stats['chunk_times'].append(chunk_time)
            self.processing_stats['processed_chunks'] = idx
            
            if progress_callback:
                progress_callback(idx, len(scenarios))
        
        # Spakuj do ZIP
        zip_path = results_dir / f"automation_tests_{task_id}.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for java_file in generated_files:
                zipf.write(java_file, java_file.name)
        
        print(f"[AUTOMATYZACJA] Wygenerowano {len(generated_files)} plików testowych -> {zip_path.name}")
        
        return zip_path
    
    def _load_scenarios_from_excel(self, excel_path: Path) -> List[Dict]:
        """Wczytuje scenariusze z pliku Excel."""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(excel_path, data_only=True)
            ws = wb.active
            
            scenarios = {}  # Grupowanie po test_case_id
            
            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row[0]:  # Pusta linia
                    continue
                    
                test_case_id = str(row[0]) if row[0] else ''
                path_type = str(row[1]) if row[1] else ''
                test_path = str(row[2]) if row[2] else ''
                scenario_title = str(row[3]) if row[3] else ''
                step_number = row[4] if row[4] else 0
                action = str(row[5]) if row[5] else ''
                expected_result = str(row[6]) if row[6] else ''
                prerequisites = str(row[7]) if row[7] else ''
                doc_section = str(row[8]) if len(row) > 8 and row[8] else ''
                
                if test_case_id not in scenarios:
                    scenarios[test_case_id] = {
                        'test_case_id': test_case_id,
                        'path_type': path_type,
                        'test_path': test_path,
                        'title': scenario_title,
                        'prerequisites': prerequisites,
                        'doc_section': doc_section,
                        'steps': []
                    }
                
                scenarios[test_case_id]['steps'].append({
                    'step_number': step_number,
                    'action': action,
                    'expected_result': expected_result
                })
            
            return list(scenarios.values())
            
        except Exception as e:
            print(f"Błąd wczytywania Excel: {e}")
            return []
    
    def _format_scenario_for_automation(self, scenario: Dict) -> str:
        """Formatuje scenariusz do promptu automatyzacji."""
        lines = [
            f"TYTUŁ SCENARIUSZA: {scenario.get('title', '')}",
            f"ID: {scenario.get('test_case_id', '')}",
            f"TYP: {scenario.get('path_type', '')}",
            f"WYMAGANIA WSTĘPNE: {scenario.get('prerequisites', 'Brak')}",
            "",
            "KROKI TESTOWE:"
        ]
        
        for step in scenario.get('steps', []):
            lines.append(f"  Krok {step.get('step_number', '?')}: {step.get('action', '')}")
            lines.append(f"    Oczekiwany rezultat: {step.get('expected_result', '')}")
        
        return '\n'.join(lines)
    
    def _extract_java_code(self, response: str, scenario: Dict) -> str:
        """Wyodrębnia kod Java z odpowiedzi LLM."""
        import re
        
        # Spróbuj wyodrębnić blok kodu Java
        java_match = re.search(r'```java\s*(.*?)\s*```', response, re.DOTALL)
        if java_match:
            return java_match.group(1).strip()
        
        # Spróbuj wyodrębnić dowolny blok kodu
        code_match = re.search(r'```\s*(.*?)\s*```', response, re.DOTALL)
        if code_match:
            return code_match.group(1).strip()
        
        # Jeśli brak bloków kodu, zwróć całą odpowiedź
        # ale usuń ewentualne wyjaśnienia na początku/końcu
        lines = response.strip().split('\n')
        code_lines = []
        in_code = False
        
        for line in lines:
            if 'package ' in line or 'import ' in line or 'public class' in line:
                in_code = True
            if in_code:
                code_lines.append(line)
        
        if code_lines:
            return '\n'.join(code_lines)
        
        # Fallback - wygeneruj podstawowy szablon
        return self._generate_fallback_test_class(scenario)
    
    def _generate_class_name(self, title: str) -> str:
        """Generuje nazwę klasy Java z tytułu scenariusza."""
        import re
        
        # Usuń polskie znaki
        replacements = {
            'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n',
            'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
            'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
            'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
        }
        
        for pl, en in replacements.items():
            title = title.replace(pl, en)
        
        # Usuń znaki specjalne
        title = re.sub(r'[^a-zA-Z0-9\s]', '', title)
        
        # Konwertuj na CamelCase
        words = title.split()
        class_name = ''.join(word.capitalize() for word in words)
        
        # Dodaj suffix Test
        if not class_name.endswith('Test'):
            class_name += 'Test'
        
        # Upewnij się że nie zaczyna się od cyfry
        if class_name and class_name[0].isdigit():
            class_name = 'Test' + class_name
        
        return class_name or 'GeneratedTest'
    
    def _generate_fallback_test_class(self, scenario: Dict) -> str:
        """Generuje podstawowy szablon klasy testowej gdy LLM zawiedzie."""
        class_name = self._generate_class_name(scenario.get('title', 'Test'))
        steps = scenario.get('steps', [])
        
        step_methods = []
        step_calls = []
        
        for step in steps:
            step_num = step.get('step_number', 1)
            action = step.get('action', 'Wykonaj akcję')
            expected = step.get('expected_result', 'Zweryfikuj wynik')
            
            method_name = f"krok{step_num}"
            step_methods.append(f'''
    @Step("Krok {step_num}: {action[:50]}")
    private void {method_name}() {{
        log.info("Krok {step_num}: {action[:80]} | Selektor: SELEKTOR_DO_UZUPELNIENIA");
        $("SELEKTOR_DO_UZUPELNIENIA").click(); // TODO: Uzupełnij akcję
        // Oczekiwany rezultat: {expected[:80]}
    }}''')
            step_calls.append(f"        {method_name}();")
        
        return f'''package tests;

import com.codeborne.selenide.Selenide;
import io.qameta.allure.*;
import lombok.extern.slf4j.Slf4j;
import org.testng.annotations.*;

import static com.codeborne.selenide.Selenide.*;

@Slf4j
@Feature("{scenario.get('doc_section', 'Testy automatyczne')}")
@Story("{scenario.get('title', 'Test automatyczny')}")
public class {class_name} {{

    @BeforeMethod
    public void setUp() {{
        // TODO: Konfiguracja WebDrivera
        log.info("Rozpoczynam test: {scenario.get('title', '')}");
    }}

    @AfterMethod
    public void tearDown() {{
        Selenide.closeWebDriver();
    }}

    @Test
    @Description("{scenario.get('title', 'Test automatyczny')}")
    public void test{class_name.replace('Test', '')}() {{
        log.info("Wymagania wstępne: {scenario.get('prerequisites', 'Brak')}");
{chr(10).join(step_calls)}
    }}
{''.join(step_methods)}
}}'''
    
    def _get_default_automation_prompt(self) -> str:
        """Domyślny prompt dla automatyzacji (fallback)."""
        return """Wygeneruj klasę testową w Javie z użyciem Selenium, Selenide, TestNG i Allure.
Każdy krok scenariusza powinien być osobną metodą z adnotacją @Step.
W miejscu selektora wstaw placeholder "SELEKTOR_DO_UZUPELNIENIA".
Zwróć tylko kod Java, bez dodatkowych komentarzy."""
